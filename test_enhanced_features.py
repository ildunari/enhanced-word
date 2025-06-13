#!/usr/bin/env python3
"""
Test script for enhanced Word MCP server features.
Tests the new Tier 1 features: Enhanced Search/Replace, Review Tools, and Section Management.
"""

import asyncio
import os
from pathlib import Path
from docx import Document

# Import our enhanced tools
from word_document_server.tools.content_tools import (
    enhanced_search_and_replace, 
    format_specific_words,
    format_research_paper_terms
)
from word_document_server.tools.review_tools import (
    manage_comments,
    extract_track_changes,
    generate_review_summary,
)
from word_document_server.tools.section_tools import (
    get_sections,
    generate_table_of_contents,
)

# Path to the generated test document
# NOTE: Each test function creates its own test document to avoid race conditions


def create_test_document():
    """Create a test document for demonstrating enhanced features."""
    import time
    timestamp = int(time.time() * 1000)  # Millisecond timestamp
    filename = f"test_enhanced_features_{timestamp}.docx"
    
    # Create a document with academic content
    doc = Document()
    
    # Add title
    title = doc.add_heading('PCL Mesophase Drug Delivery Research', 0)
    
    # Add abstract section
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This study investigates the use of polycaprolactone (PCL) mesophases for controlled '
        'drug delivery of three compounds: dolutegravir (DTG), meloxicam (MLX), and '
        'dexamethasone (DEX). Statistical analysis revealed significant correlations '
        '(p < 0.05) between mesophase content and release kinetics at both 25°C and 50°C.'
    )
    
    # Add introduction section
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'Polycaprolactone has emerged as a promising biodegradable polymer for pharmaceutical '
        'applications. The crystallinity of PCL can be modulated through thermomechanical '
        'processing to create mesophase structures that influence drug release profiles.'
    )
    
    doc.add_heading('Drug Compounds', level=2)
    doc.add_paragraph(
        'Three model drugs were selected: dolutegravir for HIV treatment, '
        'meloxicam as an anti-inflammatory agent, and dexamethasone as a corticosteroid. '
        'Each compound exhibits different solubility characteristics affecting release kinetics.'
    )
    
    # Add methods section
    doc.add_heading('Methods', level=1)
    doc.add_paragraph(
        'Samples were processed using compression molding at temperatures of 25°C and 50°C. '
        'X-ray diffraction analysis was performed to quantify crystallinity changes. '
        'Release studies were conducted in phosphate-buffered saline with ANOVA statistical analysis.'
    )
    
    # Add results section
    doc.add_heading('Results', level=1)
    doc.add_paragraph(
        'Significant differences were observed between treatment groups. '
        'The correlation between mesophase content and drug release was highly significant '
        'with r² values exceeding 0.85 for all compounds tested.'
    )
    
    # Add table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Add table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Compound'
    hdr_cells[1].text = 'Mesophase %'
    hdr_cells[2].text = 'Release Rate (μg/mL/h)'
    
    # Add data
    data = [
        ['DTG', '23.5 ± 2.1', '15.2 ± 1.8'],
        ['MLX', '31.7 ± 3.4', '22.1 ± 2.5'],
        ['DEX', '18.9 ± 1.9', '12.7 ± 1.4']
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Add conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'This research demonstrates the potential of PCL mesophase engineering for '
        'precise control of drug release kinetics. The significant correlations observed '
        'support the hypothesis that processing-induced mesophases can be leveraged '
        'for pharmaceutical applications.'
    )
    
    doc.save(filename)
    return filename


def test_enhanced_search_replace():
    """Test the enhanced search and replace functionality."""
    print("\n=== TESTING ENHANCED SEARCH AND REPLACE ===")

    # Create test document for this test
    filename = create_test_document()
    
    # Test 1: Basic enhanced search and replace with formatting
    print("Test 1: Replace 'PCL' with formatted version...")
    result = enhanced_search_and_replace(
        filename=filename,
        find_text="PCL",
        replace_text="PCL",
        apply_formatting=True,
        bold=True,
        color="blue"
    )
    print(f"Result: {result}")
    
    # Test 2: Format specific research terms
    print("\nTest 2: Format research paper terms...")
    result = format_research_paper_terms(filename)
    print(f"Result: {result}")
    
    # Test 3: Format statistical terms
    print("\nTest 3: Format statistical significance values...")
    result = format_specific_words(
        filename=filename,
        word_list=["p < 0.05", "r²", "±"],
        bold=True,
        color="red",
        whole_words_only=False
    )
    print(f"Result: {result}")
    
    # Test 4: Delete 'mesophase' occurrences using empty replacement
    print("\nTest 4: Delete 'mesophase' occurrences...")
    result = enhanced_search_and_replace(
        filename=filename,
        find_text="mesophase",
        replace_text="",  # Empty string for deletion
        match_case=False
    )
    print(f"Result: {result}")


def test_review_tools():
    """Test the review and collaboration tools."""
    print("\n=== TESTING REVIEW TOOLS ===")
    
    # Create test document for this test
    filename = create_test_document()
    
    # Test 1: Add a comment
    print("Test 1: Adding a comment to the abstract...")
    result = manage_comments(
        filename=filename,
        action="add",
        paragraph_index=2,  # Abstract paragraph
        comment_text="Consider adding more details about the mechanism of mesophase formation.",
        author="Dr. Smith",
    )
    print(f"Result: {result}")

    # Test 2: Extract comments (if any exist)
    print("\nTest 2: Extracting comments...")
    result = manage_comments(filename=filename, action="list")
    print(f"Result: {result}")
    
    # Test 3: Extract track changes (if any exist)
    print("\nTest 3: Extracting track changes...")
    result = extract_track_changes(filename=filename)
    print(f"Result: {result}")
    
    # Test 4: Generate review summary
    print("\nTest 4: Generating review summary...")
    result = generate_review_summary(filename=filename)
    print(f"Result: {result}")


def test_section_tools():
    """Test the section management tools."""
    print("\n=== TESTING SECTION MANAGEMENT TOOLS ===")
    
    # Create test document for this test
    filename = create_test_document()
    
    # Test 1: Extract sections by heading
    print("Test 1: Extracting document sections...")
    result = asyncio.run(get_sections(filename=filename, mode="overview"))
    print(f"Result: {result}")
    
    # Test 2: Extract specific section content
    print("\nTest 2: Extracting 'Methods' section content...")
    result = asyncio.run(get_sections(filename=filename, mode="content", section_title="Methods"))
    print(f"Result: {result}")
    
    # Test 3: Generate table of contents
    print("\nTest 3: Generating table of contents...")
    result = asyncio.run(generate_table_of_contents(filename))
    print(f"Result: {result}")
    
    # Test 4: Get section overview as JSON for statistics
    print("\nTest 4: Getting section statistics...")
    stats = asyncio.run(get_sections(filename=filename, mode="overview", output_format="json"))
    print(f"Result: {stats}")



