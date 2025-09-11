"""
Citation management utilities for Word documents.

This module provides functions to extract, analyze, and manage citations
with special support for EndNote citations stored as fields. These are
plain async functions intended to be registered by the main MCP server.
"""

from typing import Optional, Dict, Any, List
import json
import os
from docx import Document

from word_document_server.utils.citation_utils import (
    extract_all_citations_from_document,
    extract_citations_from_paragraph,
    create_citation_field_xml,
    extract_fields_from_run
)
from word_document_server.utils.session_utils import resolve_document_path


async def list_citations(
    document_id: str = None,
    filename: str = None,
    include_metadata: bool = True,
    group_by: str = "paragraph"
) -> str:
    """Extract and list all citations in a Word document.
    
    This tool extracts citations that are stored as fields in Word documents,
    particularly EndNote citations. It can identify citation locations, extract
    metadata, and provide a comprehensive view of all citations in the document.
    
    Args:
        document_id (str): Session document ID (preferred)
        filename (str): Path to the Word document (legacy)
        include_metadata (bool): Whether to include detailed citation metadata
            - True: Include author, year, record numbers, EndNote XML (default)
            - False: Only basic citation text and location
        group_by (str): How to group citations in the output:
            - "paragraph": Group by paragraph (default)
            - "none": Flat list of all citations
            - "reference": Group by unique reference ID
    
    Returns:
        str: JSON containing citation information:
            - total_citations: Total number of citations found
            - unique_references: Number of unique references
            - citations: Detailed citation information based on grouping
            - summary: Statistical summary of citations
    
    Examples:
        # List all citations with full metadata
        result = await list_citations(document_id="research_paper")
        
        # Get basic citation list without metadata
        result = await list_citations(document_id="thesis", include_metadata=False)
        
        # Group citations by reference to see usage patterns
        result = await list_citations(document_id="paper", group_by="reference")
    """
    # Resolve document path
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        doc = Document(filename)
        all_citations = extract_all_citations_from_document(doc)
        
        # Format based on grouping preference
        if group_by == "paragraph":
            # Default format from extract_all_citations_from_document
            result = all_citations
            
            # Remove metadata if not requested
            if not include_metadata:
                for para in result['citations_by_paragraph']:
                    for citation in para['citations']:
                        citation.pop('metadata', None)
                        citation.pop('instruction', None)
        
        elif group_by == "none":
            # Flat list of all citations
            flat_citations = []
            for para_info in all_citations['citations_by_paragraph']:
                for citation in para_info['citations']:
                    citation_copy = citation.copy()
                    citation_copy['paragraph_index'] = para_info['paragraph_index']
                    citation_copy['paragraph_text'] = para_info['paragraph_text'][:100] + "..."
                    
                    if not include_metadata:
                        citation_copy.pop('metadata', None)
                        citation_copy.pop('instruction', None)
                    
                    flat_citations.append(citation_copy)
            
            result = {
                'total_citations': all_citations['total_citations'],
                'unique_references': all_citations['unique_references'],
                'citations': flat_citations,
                'citation_summary': all_citations['citation_summary']
            }
        
        elif group_by == "reference":
            # Group by unique reference
            refs_dict = {}
            
            for para_info in all_citations['citations_by_paragraph']:
                for citation in para_info['citations']:
                    ref_id = citation.get('metadata', {}).get('record_number', 'unknown')
                    
                    if ref_id not in refs_dict:
                        refs_dict[ref_id] = {
                            'reference_id': ref_id,
                            'occurrences': [],
                            'metadata': citation.get('metadata', {}) if include_metadata else {},
                            'total_occurrences': 0
                        }
                    
                    occurrence = {
                        'paragraph_index': para_info['paragraph_index'],
                        'character_position': citation['character_position'],
                        'display_text': citation['display_text'],
                        'context': para_info['paragraph_text'][:100] + "..."
                    }
                    
                    refs_dict[ref_id]['occurrences'].append(occurrence)
                    refs_dict[ref_id]['total_occurrences'] += 1
            
            result = {
                'total_citations': all_citations['total_citations'],
                'unique_references': all_citations['unique_references'],
                'citations_by_reference': list(refs_dict.values()),
                'citation_summary': all_citations['citation_summary']
            }
        
        else:
            return f"Invalid group_by value: {group_by}. Must be one of: paragraph, none, reference"
        
        return json.dumps(result, indent=2)
        
    except Exception as e:
        return f"Failed to extract citations: {str(e)}"


async def get_citation_at_position(
    document_id: str = None,
    filename: str = None,
    paragraph_index: int = None,
    character_position: Optional[int] = None
) -> str:
    """Get detailed information about a citation at a specific position.
    
    This tool retrieves comprehensive information about a citation at a given
    location in the document, including its metadata, surrounding context,
    and field structure.
    
    Args:
        document_id (str): Session document ID (preferred)
        filename (str): Path to the Word document (legacy)
        paragraph_index (int): Zero-based paragraph index
        character_position (int, optional): Character position within paragraph
            - If not provided, returns all citations in the paragraph
    
    Returns:
        str: JSON containing citation details:
            - paragraph_text: Full text of the paragraph
            - citations: List of citations found
            - field_structure: Raw field XML for advanced users
    
    Examples:
        # Get citation at specific position
        result = await get_citation_at_position(
            document_id="paper",
            paragraph_index=5,
            character_position=120
        )
        
        # Get all citations in a paragraph
        result = await get_citation_at_position(
            document_id="thesis",
            paragraph_index=10
        )
    """
    # Resolve document path
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    if paragraph_index is None:
        return "paragraph_index is required"
    
    try:
        doc = Document(filename)
        
        if paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index: {paragraph_index}. Document has {len(doc.paragraphs)} paragraphs"
        
        paragraph = doc.paragraphs[paragraph_index]
        citations = extract_citations_from_paragraph(paragraph)
        
        result = {
            'paragraph_index': paragraph_index,
            'paragraph_text': paragraph.text,
            'total_citations_in_paragraph': len(citations)
        }
        
        if character_position is not None:
            # Find citation at or near the position
            matching_citations = []
            for citation in citations:
                cite_start = citation['character_position']
                cite_end = cite_start + len(citation['display_text'])
                
                # Check if position is within citation
                if cite_start <= character_position <= cite_end:
                    matching_citations.append(citation)
            
            result['citations_at_position'] = matching_citations
            result['character_position'] = character_position
        else:
            result['all_citations'] = citations
        
        return json.dumps(result, indent=2)
        
    except Exception as e:
        return f"Failed to get citation at position: {str(e)}"


async def copy_existing_citation(
    document_id: str = None,
    filename: str = None,
    source_paragraph: int = None,
    citation_index: int = 0
) -> str:
    """Copy citation field data for reuse elsewhere in the document.
    
    This tool extracts the complete field structure of an existing citation
    so it can be inserted elsewhere while maintaining the proper EndNote
    linking and formatting.
    
    Args:
        document_id (str): Session document ID (preferred)
        filename (str): Path to the Word document (legacy)
        source_paragraph (int): Paragraph index containing the citation to copy
        citation_index (int): Which citation to copy if multiple in paragraph (default: 0)
    
    Returns:
        str: JSON containing copyable citation data:
            - citation_data: Complete citation information
            - field_instruction: Field instruction for insertion
            - display_text: What will be shown in the document
            - insert_instructions: How to use this data
    
    Examples:
        # Copy the first citation from paragraph 10
        result = await copy_existing_citation(
            document_id="paper",
            source_paragraph=10
        )
        
        # Copy the second citation from paragraph 15
        result = await copy_existing_citation(
            document_id="thesis",
            source_paragraph=15,
            citation_index=1
        )
    """
    # Resolve document path
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    if source_paragraph is None:
        return "source_paragraph is required"
    
    try:
        doc = Document(filename)
        
        if source_paragraph >= len(doc.paragraphs):
            return f"Invalid paragraph index: {source_paragraph}. Document has {len(doc.paragraphs)} paragraphs"
        
        paragraph = doc.paragraphs[source_paragraph]
        citations = extract_citations_from_paragraph(paragraph)
        
        if not citations:
            return f"No citations found in paragraph {source_paragraph}"
        
        if citation_index >= len(citations):
            return f"Invalid citation_index: {citation_index}. Paragraph has {len(citations)} citations"
        
        citation = citations[citation_index]
        
        # Prepare copyable citation data
        result = {
            'source_location': {
                'paragraph_index': source_paragraph,
                'citation_index': citation_index
            },
            'citation_data': {
                'display_text': citation['display_text'],
                'field_type': citation['field_type'],
                'metadata': citation['metadata'],
                'instruction': citation['instruction']
            },
            'insert_instructions': (
                "This citation data can be used with text editing tools to insert "
                "the citation elsewhere. The 'instruction' field contains the complete "
                "EndNote field data needed to maintain proper citation linking."
            ),
            'example_usage': (
                "To insert this citation in another location, use the enhanced_search_and_replace "
                "tool with preserve_fields=True, or use add_text_content with field support."
            )
        }
        
        return json.dumps(result, indent=2)
        
    except Exception as e:
        return f"Failed to copy citation: {str(e)}"


async def analyze_citation_distribution(
    document_id: str = None,
    filename: str = None,
    section_headers: Optional[List[str]] = None
) -> str:
    """Analyze the distribution and patterns of citations throughout the document.
    
    This tool provides statistical analysis of how citations are distributed
    across the document, helping identify areas that may need more references
    or sections with heavy citation density.
    
    Args:
        document_id (str): Session document ID (preferred)
        filename (str): Path to the Word document (legacy)
        section_headers (List[str], optional): List of section header patterns
            to analyze citation distribution by section
    
    Returns:
        str: JSON containing distribution analysis:
            - overall_statistics: Document-wide citation statistics
            - distribution_by_position: Citation density throughout document
            - most_cited_references: Top referenced sources
            - citation_gaps: Paragraphs without citations
            - section_analysis: Citation breakdown by section (if headers provided)
    
    Examples:
        # Basic distribution analysis
        result = await analyze_citation_distribution(document_id="thesis")
        
        # Analyze by sections
        result = await analyze_citation_distribution(
            document_id="paper",
            section_headers=["Introduction", "Methods", "Results", "Discussion"]
        )
    """
    # Resolve document path
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        doc = Document(filename)
        all_citations = extract_all_citations_from_document(doc)
        
        # Calculate overall statistics
        total_paragraphs = len(doc.paragraphs)
        paragraphs_with_citations = len(all_citations['citations_by_paragraph'])
        
        # Find citation gaps (consecutive paragraphs without citations)
        cited_paragraphs = {para['paragraph_index'] for para in all_citations['citations_by_paragraph']}
        
        gaps = []
        gap_start = None
        
        for i in range(total_paragraphs):
            if i not in cited_paragraphs:
                if gap_start is None:
                    gap_start = i
            else:
                if gap_start is not None:
                    gaps.append({
                        'start': gap_start,
                        'end': i - 1,
                        'length': i - gap_start
                    })
                    gap_start = None
        
        # Handle gap at end of document
        if gap_start is not None:
            gaps.append({
                'start': gap_start,
                'end': total_paragraphs - 1,
                'length': total_paragraphs - gap_start
            })
        
        # Find most cited references
        ref_counts = {}
        for para_info in all_citations['citations_by_paragraph']:
            for citation in para_info['citations']:
                ref_id = citation.get('metadata', {}).get('record_number', 'unknown')
                ref_counts[ref_id] = ref_counts.get(ref_id, 0) + 1
        
        most_cited = sorted(ref_counts.items(), key=lambda x: x[1], reverse=True)[:10]
        
        result = {
            'overall_statistics': {
                'total_paragraphs': total_paragraphs,
                'paragraphs_with_citations': paragraphs_with_citations,
                'paragraphs_without_citations': total_paragraphs - paragraphs_with_citations,
                'citation_coverage_percentage': (paragraphs_with_citations / total_paragraphs * 100) if total_paragraphs > 0 else 0,
                'total_citations': all_citations['total_citations'],
                'unique_references': len(all_citations['unique_references']),
                'average_citations_per_paragraph': all_citations['citation_summary']['average_citations_per_paragraph']
            },
            'most_cited_references': [
                {'reference_id': ref_id, 'citation_count': count}
                for ref_id, count in most_cited
            ],
            'citation_gaps': sorted(gaps, key=lambda x: x['length'], reverse=True)[:5],
            'distribution_summary': {
                'first_citation_paragraph': min(cited_paragraphs) if cited_paragraphs else None,
                'last_citation_paragraph': max(cited_paragraphs) if cited_paragraphs else None,
                'longest_gap_without_citations': max(gap['length'] for gap in gaps) if gaps else 0
            }
        }
        
        # Section analysis if headers provided
        if section_headers:
            section_data = {}
            current_section = "Before first section"
            section_start = 0
            
            for i, paragraph in enumerate(doc.paragraphs):
                # Check if this paragraph is a section header
                para_text = paragraph.text.strip()
                for header in section_headers:
                    if header.lower() in para_text.lower():
                        # Save previous section data
                        if current_section not in section_data:
                            section_data[current_section] = {
                                'start_paragraph': section_start,
                                'end_paragraph': i - 1,
                                'citations': 0,
                                'paragraphs': i - section_start
                            }
                        
                        current_section = header
                        section_start = i + 1
                        break
            
            # Handle last section
            section_data[current_section] = {
                'start_paragraph': section_start,
                'end_paragraph': total_paragraphs - 1,
                'citations': 0,
                'paragraphs': total_paragraphs - section_start
            }
            
            # Count citations per section
            for para_info in all_citations['citations_by_paragraph']:
                para_idx = para_info['paragraph_index']
                
                for section, data in section_data.items():
                    if data['start_paragraph'] <= para_idx <= data['end_paragraph']:
                        data['citations'] += len(para_info['citations'])
                        break
            
            # Calculate section statistics
            section_analysis = []
            for section, data in section_data.items():
                if data['paragraphs'] > 0:
                    section_analysis.append({
                        'section': section,
                        'paragraph_range': f"{data['start_paragraph']}-{data['end_paragraph']}",
                        'total_paragraphs': data['paragraphs'],
                        'total_citations': data['citations'],
                        'citations_per_paragraph': data['citations'] / data['paragraphs']
                    })
            
            result['section_analysis'] = section_analysis
        
        return json.dumps(result, indent=2)
        
    except Exception as e:
        return f"Failed to analyze citation distribution: {str(e)}"


async def citations(
    action: str,
    document_id: str = None,
    filename: str = None,
    # list params
    include_metadata: bool = True,
    group_by: str = "paragraph",
    # at_position params
    paragraph_index: Optional[int] = None,
    character_position: Optional[int] = None,
    # copy params
    source_paragraph: Optional[int] = None,
    citation_index: int = 0,
    # distribution params
    section_headers: Optional[List[str]] = None,
) -> str:
    """Consolidated citation tool.

    action:
      - "list": list all citations (supports include_metadata, group_by)
      - "at_position": get citation(s) at a paragraph/character position
      - "copy": copy an existing citation’s field data from a paragraph
      - "distribution": analyze document-wide citation distribution
    """
    valid = {"list", "at_position", "copy", "distribution"}
    if action not in valid:
        return f"Invalid action: {action}. Must be one of: {', '.join(sorted(valid))}"

    if action == "list":
        return await list_citations(
            document_id=document_id,
            filename=filename,
            include_metadata=include_metadata,
            group_by=group_by,
        )
    if action == "at_position":
        return await get_citation_at_position(
            document_id=document_id,
            filename=filename,
            paragraph_index=paragraph_index,
            character_position=character_position,
        )
    if action == "copy":
        return await copy_existing_citation(
            document_id=document_id,
            filename=filename,
            source_paragraph=source_paragraph,
            citation_index=citation_index,
        )
    if action == "distribution":
        return await analyze_citation_distribution(
            document_id=document_id,
            filename=filename,
            section_headers=section_headers,
        )

    return f"Invalid action: {action}"
