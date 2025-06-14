"""Equation insertion tool for Enhanced Word MCP server.

Adds editable Word equations by converting LaTeX to OMML and injecting the
XML fragment into the document.  Falls back to an error message if the
necessary external tools are missing.
"""

from __future__ import annotations

from typing import Optional

from docx import Document
from docx.oxml import OxmlElement, parse_xml

from word_document_server.utils.session_utils import resolve_document_path
from word_document_server.utils.equation_utils import latex_to_omml
from word_document_server.utils.file_utils import check_file_writeable


async def insert_equation(
    document_id: str = None,
    filename: str = None,
    latex: str | None = None,
    display_style: bool = False,
    position: str = "end",
    paragraph_index: Optional[int] = None,
) -> str:
    """Insert an editable equation converted from LaTeX.

    Parameters
    ----------
    document_id / filename
        Identify the document (session preferred).
    latex : str
        LaTeX math expression, e.g. ``E = mc^2`` (without ``$`` delimiters).
    display_style : bool
        If True, insert as a separate paragraph (block equation).  If False,
        insert inline.
    position : str
        Where to insert: "end" | "before" | "after" | "beginning".
    paragraph_index : int, optional
        Target paragraph for before/after modes.
    """
    # Validate params
    if not latex or not latex.strip():
        return "Error: 'latex' parameter is required"

    valid_positions = {"end", "before", "after", "beginning"}
    if position not in valid_positions:
        return f"Invalid position: {position}. Must be one of: {', '.join(valid_positions)}"

    if position in {"before", "after"} and paragraph_index is None:
        return "paragraph_index parameter is required for position 'before' or 'after'"

    # Resolve path
    file_path, error = resolve_document_path(document_id, filename)
    if error:
        return error

    # Check writeable and trigger undo snapshot
    ok, err_msg = check_file_writeable(file_path)
    if not ok:
        return err_msg

    # Convert LaTeX → OMML
    success, omml_or_err = latex_to_omml(latex)
    if not success:
        return omml_or_err
    omml_xml = omml_or_err

    try:
        doc = Document(file_path)

        total_paras = len(doc.paragraphs)

        # Validate index bounds when needed
        if position in {"before", "after"}:
            if paragraph_index < 0 or paragraph_index >= total_paras:
                return f"paragraph_index out of range: 0-{total_paras-1}"

        # Decide which paragraph will host the equation
        if display_style or position in {"before", "after"}:
            # We create a brand-new paragraph; later we will move it if needed
            p = doc.add_paragraph()
        else:
            if position == "end":
                p = doc.paragraphs[-1]
            elif position == "beginning":
                p = doc.paragraphs[0]
            else:  # Should not hit
                p = doc.paragraphs[-1]

        # For inline, ensure we have a run to host equation
        run = p.add_run()

        # Inject OMML fragment
        omml_elem = parse_xml(omml_xml)
        run._r.append(omml_elem)

        # Move paragraph to correct location if we created a new one
        if display_style or position in {"before", "after"}:
            if position == "beginning":
                doc._body._element.insert(0, p._element)
            elif position == "before":
                target_p = doc.paragraphs[paragraph_index]
                target_p._element.addprevious(p._element)
            elif position == "after":
                target_p = doc.paragraphs[paragraph_index]
                target_p._element.addnext(p._element)

        # Ensure math namespace present in document root
        from docx.oxml.ns import qn  # local import to avoid global dependency

        root = doc._element
        if "m" not in root.nsmap:
            root.set(qn("xmlns:m"), "http://schemas.openxmlformats.org/officeDocument/2006/math")

        doc.save(file_path)
        return "Equation inserted successfully"
    except Exception as e:
        return f"Failed to insert equation: {e}"


__all__ = ["insert_equation"] 