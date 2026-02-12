"""
Content tools for Word Document Server.

These tools add various types of content to Word documents,
including headings, paragraphs, tables, images, and page breaks.
"""
import os
import re
from typing import List, Optional
from docx import Document
from docx.shared import Inches, Pt

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension, validate_docx_path, sanitize_file_path
from word_document_server.utils.document_utils import find_and_replace_text
from word_document_server.utils.session_utils import resolve_document_path
from word_document_server.core.styles import ensure_heading_style, ensure_table_style
from word_document_server.utils.equation_utils import latex_to_omml
from word_document_server.utils.citation_utils import extract_fields_from_run, format_run_with_citation_awareness
from word_document_server.utils.limits import (
    LIMIT_EXCEEDED,
    REGEX_COMPLEXITY_BLOCKED,
    check_doc_size_for_operation,
    get_max_matches_per_call,
    get_max_regex_pattern_chars,
    get_max_regex_scan_chars,
    get_max_search_output_chars,
    get_regex_timeout_ms,
    is_regex_timeout_supported,
)
from docx.oxml import parse_xml


def _truncate_message(message: str, max_chars: int) -> str:
    """Prevent oversized textual tool responses."""
    if len(message) <= max_chars:
        return message
    suffix = f"... [truncated to {max_chars} chars]"
    keep = max_chars - len(suffix)
    if keep <= 0:
        return suffix[-max_chars:]
    return message[:keep] + suffix


def _regex_complexity_reason(pattern: str) -> Optional[str]:
    """Return a reason string when a pattern looks pathologically expensive."""
    nested_quantifier = re.search(
        r"\((?:[^()\\]|\\.)*(?:\*|\+|\{\d+,?\d*\})(?:[^()\\]|\\.)*\)\s*(?:\*|\+|\{\d+,?\d*\})",
        pattern,
    )
    if nested_quantifier:
        return "nested quantifiers in grouped expression"

    wildcard_repetition = re.search(r"\((?:\.\*|\.\+|\\w\+|\\s\+)\)\+", pattern)
    if wildcard_repetition:
        return "repeated wildcard-like capture group"

    alternations = pattern.count("|")
    if alternations > 256:
        return f"alternation count {alternations} exceeds safety threshold 256"

    return None


async def add_text_content(
    document_id: str = None,
    filename: str = None,
    text: str = None,
    content_type: str = "paragraph",
    level: Optional[int] = None,
    style: Optional[str] = None,
    position: str = "end",
    insert_before_paragraph: Optional[int] = None,
    insert_after_paragraph: Optional[int] = None
) -> str:
    """Unified text content addition function for comprehensive document content management.
    
    This function consolidates paragraph and heading creation into a single comprehensive tool.
    It replaces add_paragraph and add_heading with enhanced positioning, styling, and
    formatting options for professional document creation and editing workflows.
    
    Args:
        document_id (str): Session document ID (preferred)
        filename (str): Path to the Word document (legacy, for backward compatibility)
        
        text (str): Content text to add to the document
            - Supports plain text with basic formatting preservation
            - Can include special characters and unicode
            - Line breaks will be preserved in paragraph content
            - Example: "This is the introduction paragraph with important concepts."
        
        content_type (str): Type of content element to create:
            - "paragraph": Regular text paragraph (default)
            - "heading": Document heading/title with hierarchical level
            - Paragraphs are for body content, headings for structure
        
        level (int, optional): Heading hierarchy level (required for content_type="heading")
            - 1: Main title/chapter heading (largest)
            - 2: Section heading
            - 3: Subsection heading
            - 4-9: Sub-subsection headings (progressively smaller)
            - Only used when content_type="heading"
        
        style (str, optional): Document style name to apply to content
            - None: Use default paragraph/heading style (default)
            - "Normal": Standard paragraph style
            - "Quote": Indented quotation style
            - "Emphasis": Emphasized text style
            - "Caption": Figure/table caption style
            - Custom: Any style defined in document template
        
        position (str): Placement position within document:
            - "end": Append to end of document (default)
            - "beginning": Insert at document start
            - "before": Insert before specific paragraph (requires insert_before_paragraph)
            - "after": Insert after specific paragraph (requires insert_after_paragraph)
        
        insert_before_paragraph (int, optional): Zero-based paragraph index for position="before"
            - Content will be inserted before this paragraph
            - Existing paragraph indices will shift down
            - Example: 5 inserts before the 6th paragraph
        
        insert_after_paragraph (int, optional): Zero-based paragraph index for position="after"
            - Content will be inserted after this paragraph
            - Subsequent paragraph indices will shift down
            - Example: 3 inserts after the 4th paragraph
    
    Returns:
        str: Status message indicating success or failure:
            - Success: "Successfully added {content_type} at {position}"
            - Error: Specific error message with troubleshooting guidance
    
    Use Cases:
        📝 Document Creation: Build structured documents with headings and content
        📚 Academic Writing: Add chapters, sections, and content paragraphs
        📄 Report Writing: Insert executive summaries, findings, conclusions
        📋 Technical Documentation: Add procedure steps and explanations
        ✏️ Manuscript Editing: Insert new content at specific locations
        📊 Business Documents: Add formatted content with professional styling
    
    Examples:
        # Add simple paragraph at document end
        result = await add_text_content(document_id="report", 
                                       text="This paragraph summarizes the key findings of our research.")
        # Returns: "Successfully added paragraph at end"
        
        # Create main chapter heading
        result = await add_text_content(document_id="thesis", text="Chapter 3: Methodology", 
                                       content_type="heading", level=1)
        # Returns: "Successfully added heading at end"
        
        # Insert section heading with custom style
        result = await add_text_content(document_id="manuscript", text="3.1 Data Collection",
                                       content_type="heading", level=2, style="Heading 2")
        # Returns: "Successfully added heading at end"
        
        # Insert paragraph before specific location
        result = await add_text_content(document_id="document", 
                                       text="This new paragraph provides important context.",
                                       position="before", insert_before_paragraph=5)
        # Returns: "Successfully added paragraph at before"
        
        # Add quotation with quote style
        result = await add_text_content(document_id="essay", 
                                       text="To be or not to be, that is the question.",
                                       style="Quote", position="after", insert_after_paragraph=2)
        # Returns: "Successfully added paragraph at after"
        
        # Insert introduction paragraph at document beginning
        result = await add_text_content(document_id="paper", 
                                       text="This document presents comprehensive analysis of market trends.",
                                       position="beginning")
        # Returns: "Successfully added paragraph at beginning"
        
        # Add subsection heading in middle of document
        result = await add_text_content(document_id="manual", text="2.3.1 Installation Steps",
                                       content_type="heading", level=3,
                                       position="before", insert_before_paragraph=15)
        # Returns: "Successfully added heading at before"
        
        # Add emphasized conclusion paragraph
        result = await add_text_content(document_id="summary",
                                       text="In conclusion, the results demonstrate significant improvements.",
                                       style="Emphasis", position="end")
        # Returns: "Successfully added paragraph at end"
    
    Error Handling:
        - Document not found: "Document '{document_id}' not found in sessions"
        - File not writable: "Cannot modify document: {reason}. Consider creating a copy first."
        - Invalid content_type: "Invalid content_type: {type}. Must be one of: paragraph, heading"
        - Invalid position: "Invalid position: {pos}. Must be one of: end, beginning, before, after"
        - Missing level for heading: "Heading level (1-9) is required for content_type='heading'"
        - Invalid level: "Invalid heading level: {level}. Must be between 1-9"
        - Missing position parameter: "insert_before_paragraph required for position='before'"
        - Invalid paragraph index: "Paragraph index {index} is out of range"
        - Empty text: "Text content cannot be empty"
        - Style not found: "Style '{style}' not found in document template"
        - Document corruption: "Error adding content: {error_details}"
    
    Document Structure Workflow:
        1. Planning: Design document hierarchy with heading levels
        2. Structure Creation: Add main headings (level 1) first
        3. Section Development: Add subsection headings (levels 2-3)
        4. Content Addition: Insert paragraphs under appropriate headings
        5. Refinement: Use positioning options to reorganize content
        6. Styling: Apply consistent styles throughout document
    
    Academic Writing Best Practices:
        - Use level 1 for chapter headings
        - Use level 2 for major section headings
        - Use level 3 for subsection headings
        - Maintain consistent heading hierarchy
        - Insert content paragraphs after relevant headings
        - Use appropriate styles for different content types
    
    Performance Notes:
        - Large documents may take longer for content insertion
        - Position-specific insertions require document traversal
        - Style application adds minimal processing time
        - Batch multiple content additions for efficiency
        - Complex positioning operations may be slower than simple appends
    
    Integration with Other Tools:
        - Use get_sections to understand document structure before insertion
        - Use get_text to verify content placement after insertion
        - Combine with add_note for comprehensive content creation
        - Use manage_protection to control editing permissions
    """
    from word_document_server.utils.session_utils import resolve_document_path
    
    # Resolve document path from document_id or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg

    # Validate required parameters
    if not text:
        return "Error: text parameter is required"
    
    # Validate content_type parameter
    valid_types = ["paragraph", "heading"]
    if content_type not in valid_types:
        return f"Invalid content_type: {content_type}. Must be one of: {', '.join(valid_types)}"
    
    # Validate position parameter
    valid_positions = ["end", "beginning", "before", "after"]
    if position not in valid_positions:
        return f"Invalid position: {position}. Must be one of: {', '.join(valid_positions)}"
    
    # Validate heading-specific parameters
    if content_type == "heading":
        if level is None:
            return "Invalid parameter: level is required when content_type='heading'"
        
        try:
            level = int(level)
            if level < 1 or level > 9:
                return f"Invalid heading level: {level}. Level must be between 1 and 9."
        except (ValueError, TypeError):
            return "Invalid parameter: level must be an integer between 1 and 9"
    
    # Validate position-specific parameters
    if position == "before" and insert_before_paragraph is None:
        return "Invalid parameter: insert_before_paragraph is required when position='before'"
    
    if position == "after" and insert_after_paragraph is None:
        return "Invalid parameter: insert_after_paragraph is required when position='after'"
    
    # Validate paragraph index parameters
    for param_name, param_value in [("insert_before_paragraph", insert_before_paragraph), 
                                   ("insert_after_paragraph", insert_after_paragraph)]:
        if param_value is not None:
            try:
                param_value = int(param_value)
                if param_value < 0:
                    return f"Invalid parameter: {param_name} must be a non-negative integer"
            except (ValueError, TypeError):
                return f"Invalid parameter: {param_name} must be an integer"
    
    if not text.strip():
        return "Invalid parameter: text cannot be empty"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    size_ok, size_error = check_doc_size_for_operation(filename, "enhanced_search_and_replace")
    if not size_ok:
        return size_error

    size_ok, size_error = check_doc_size_for_operation(filename, "add_text_content")
    if not size_ok:
        return size_error
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph indices against document
        if insert_before_paragraph is not None:
            insert_before_paragraph = int(insert_before_paragraph)
            if insert_before_paragraph >= len(doc.paragraphs):
                return f"Invalid insert_before_paragraph: {insert_before_paragraph}. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        if insert_after_paragraph is not None:
            insert_after_paragraph = int(insert_after_paragraph)
            if insert_after_paragraph >= len(doc.paragraphs):
                return f"Invalid insert_after_paragraph: {insert_after_paragraph}. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        # Create the content element
        if content_type == "heading":
            # Ensure heading styles exist
            ensure_heading_style(doc)
            
            try:
                if position == "end":
                    # Create the heading at end (default behavior)
                    heading = doc.add_heading(text, level=level)
                    created_element = heading
                    success_message = f"Heading '{text}' (level {level}) added"
                elif position == "beginning":
                    # Create the heading and move to beginning
                    heading = doc.add_heading(text, level=level)
                    created_element = heading
                    success_message = f"Heading '{text}' (level {level}) added"
                    # Move heading element to the beginning (it's automatically removed from current position)
                    body = doc._body._element
                    body.remove(heading._element)
                    body.insert(0, heading._element)
                else:
                    # Insert at specific position
                    heading = doc.add_heading(text, level=level)
                    created_element = heading
                    success_message = f"Heading '{text}' (level {level}) inserted"
                    
            except Exception:
                # Fallback to direct formatting if heading styles fail
                paragraph = doc.add_paragraph(text)
                paragraph.style = doc.styles['Normal']
                run = paragraph.runs[0]
                run.bold = True
                
                # Adjust size based on heading level
                if level == 1:
                    run.font.size = Pt(16)
                elif level == 2:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)
                    
                created_element = paragraph
                success_message = f"Heading '{text}' added with direct formatting"
        
        else:  # paragraph
            paragraph = doc.add_paragraph(text)
            created_element = paragraph
            success_message = f"Paragraph added"
            
            # Apply style if specified
            if style:
                try:
                    paragraph.style = style
                except KeyError:
                    paragraph.style = doc.styles['Normal']
                    success_message += f" (style '{style}' not found, used default)"
        
        # Handle positioning for before/after insertions
        if position == "before":
            # Move element to before specified paragraph
            target_paragraph = doc.paragraphs[insert_before_paragraph]
            target_paragraph._element.getparent().insert(
                list(target_paragraph._element.getparent()).index(target_paragraph._element),
                created_element._element
            )
            success_message += f" before paragraph {insert_before_paragraph}"
        
        elif position == "after":
            # Move element to after specified paragraph
            target_paragraph = doc.paragraphs[insert_after_paragraph]
            target_paragraph._element.getparent().insert(
                list(target_paragraph._element.getparent()).index(target_paragraph._element) + 1,
                created_element._element
            )
            success_message += f" after paragraph {insert_after_paragraph}"
        
        elif position == "beginning" and content_type == "paragraph":
            # Move paragraph to beginning
            doc._body._element.insert(0, created_element._element)
            success_message += " at document beginning"
        
        doc.save(filename)
        return f"{success_message} to {filename}"
    
    except Exception as e:
        return f"Failed to add {content_type}: {str(e)}"


async def add_table(document_id: str = None, filename: str = None, rows: int = None, cols: int = None, data: Optional[List[List[str]]] = None) -> str:
    """Add a table to a Word document.
    
    Args:
        document_id (str, optional): Session document identifier (preferred)
        filename (str, optional): Path to the Word document
        rows: Number of rows in the table
        cols: Number of columns in the table
        data: Optional 2D array of data to fill the table
    """
    # Resolve document path from session or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    size_ok, size_error = check_doc_size_for_operation(filename, "add_table")
    if not size_ok:
        return size_error

    # Validate rows and cols parameters before proceeding
    if rows is None or cols is None:
        return "Invalid parameters: 'rows' and 'cols' are required"
    try:
        rows = int(rows)
        cols = int(cols)
    except (ValueError, TypeError):
        return "Invalid parameters: 'rows' and 'cols' must be integers"
    if rows <= 0 or cols <= 0:
        return "Invalid parameters: 'rows' and 'cols' must be positive integers"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        table = doc.add_table(rows=rows, cols=cols)
        
        # Try to set the table style
        try:
            table.style = 'Table Grid'
        except KeyError:
            # If style doesn't exist, add basic borders
            pass
        
        # Fill table with data if provided
        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                for j, cell_text in enumerate(row_data):
                    if j >= cols:
                        break
                    table.cell(i, j).text = str(cell_text)
        
        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"
    except Exception as e:
        return f"Failed to add table: {str(e)}"


async def add_picture(document_id: str = None, filename: str = None, image_path: str = None, width: Optional[float] = None) -> str:
    """Add an image to a Word document.
    
    Args:
        document_id (str, optional): Session document identifier (preferred)
        filename (str, optional): Path to the Word document
        image_path: Path to the image file
        width: Optional width in inches (proportional scaling)
    """
    # Resolve document path from session or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    filename = ensure_docx_extension(filename)
    
    # Validate document existence
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Sanitize and validate image path
    if not image_path:
        return "Image file path is required"
    ok_img, safe_img_path, img_err = sanitize_file_path(
        image_path,
        allowed_extensions=[
            '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tif', '.tiff'
        ]
    )
    if not ok_img:
        return f"Invalid image path: {img_err}"

    # Get absolute paths for better diagnostics
    abs_filename = os.path.abspath(filename)
    abs_image_path = os.path.abspath(safe_img_path)
    
    # Validate image existence with improved error message
    if not os.path.exists(abs_image_path):
        return f"Image file not found: {abs_image_path}"
    
    # Check image file size
    try:
        image_size = os.path.getsize(abs_image_path) / 1024  # Size in KB
        if image_size <= 0:
            return f"Image file appears to be empty: {abs_image_path} (0 KB)"
    except Exception as size_error:
        return f"Error checking image file: {str(size_error)}"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(abs_filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(abs_filename)
        # Additional diagnostic info
        diagnostic = f"Attempting to add image ({abs_image_path}, {image_size:.2f} KB) to document ({abs_filename})"
        
        try:
            if width:
                doc.add_picture(abs_image_path, width=Inches(width))
            else:
                doc.add_picture(abs_image_path)
            doc.save(abs_filename)
            return f"Picture {image_path} added to {filename}"
        except Exception as inner_error:
            # More detailed error for the specific operation
            error_type = type(inner_error).__name__
            error_msg = str(inner_error)
            return f"Failed to add picture: {error_type} - {error_msg or 'No error details available'}\nDiagnostic info: {diagnostic}"
    except Exception as outer_error:
        # Fallback error handling
        error_type = type(outer_error).__name__
        error_msg = str(outer_error)
        return f"Document processing error: {error_type} - {error_msg or 'No error details available'}"


def enhanced_search_and_replace(document_id: str = None, filename: str = None, 
                                    find_text: str = None, replace_text: str = None,
                                    apply_formatting: bool = False,
                                    bold: Optional[bool] = None, 
                                    italic: Optional[bool] = None,
                                    underline: Optional[bool] = None, 
                                    color: Optional[str] = None,
                                    font_size: Optional[int] = None, 
                                    font_name: Optional[str] = None,
                                    match_case: bool = True,
                                    whole_words_only: bool = False,
                                    use_regex: bool = False,
                                    start_paragraph: Optional[int] = None,
                                    end_paragraph: Optional[int] = None,
                                    paragraph_indices: Optional[List[int]] = None,
                                    occurrence_index: Optional[int] = None,
                                    char_start: Optional[int] = None,
                                    char_end: Optional[int] = None,
                                    replace_with_equation: bool = False,
                                    latex_equation: Optional[str] = None,
                                    preserve_fields: bool = True) -> str:
    """Enhanced search and replace with formatting options, regex support, and case-insensitive matching.
    
    Provides powerful text replacement capabilities with:
    - Regex pattern matching
    - Case-sensitive/insensitive search
    - Whole word matching
    - Advanced formatting application to replaced text
    - Table content support
    
    Args:
        document_id: Session document ID (preferred)
        filename: Path to the Word document (legacy, for backward compatibility)
        find_text: Text or regex pattern to search for
        replace_text: Text to replace with (supports regex groups like $1, $2 if use_regex=True)
        apply_formatting: Whether to apply formatting to the replaced text
        bold: Set replaced text bold (True/False)
        italic: Set replaced text italic (True/False)
        underline: Set replaced text underlined (True/False)
        color: Text color for replaced text (e.g., 'red', 'blue', '#FF0000')
        font_size: Font size in points for replaced text
        font_name: Font name/family for replaced text
        match_case: Whether to match case exactly (default True, ignored if use_regex=True)
        whole_words_only: Whether to match whole words only (default False)
        use_regex: Enable regex pattern matching (default False)
        start_paragraph: Start paragraph index for range filtering
        end_paragraph: End paragraph index for range filtering
        paragraph_indices: List of paragraph indices to filter
        occurrence_index: Target occurrence index for filtering
        char_start: Start character index for range filtering
        char_end: End character index for range filtering
        replace_with_equation: Replace matched text with a LaTeX equation (default False)
        latex_equation: LaTeX equation to insert (required if replace_with_equation=True)
        preserve_fields: Preserve Word fields like EndNote citations when replacing text (default True)
    
    Returns:
        Status message with replacement count and details
        
    Examples:
        # Simple replacement
        enhanced_search_and_replace(document_id="main", find_text="old text", replace_text="new text")
        
        # Case-insensitive replacement with formatting
        enhanced_search_and_replace(document_id="main", find_text="Important", replace_text="CRITICAL", 
                                   match_case=False, apply_formatting=True, 
                                   bold=True, color="red")
        
        # Regex pattern replacement
        enhanced_search_and_replace(document_id="main", find_text=r"(\\\\d{4})-(\\\\d{2})-(\\\\d{2})", 
                                   replace_text=r"$2/$3/$1", use_regex=True)
                                   
        # Whole word replacement with font styling
        enhanced_search_and_replace(document_id="main", find_text="AI", replace_text="Artificial Intelligence", 
                                   whole_words_only=True, apply_formatting=True,
                                   font_name="Arial", font_size=12)
        
        # Replace text with equation
        enhanced_search_and_replace(document_id="main", find_text="E=mc2", 
                                   replace_with_equation=True, latex_equation="E = mc^2")
                                   
        # Insert equation at specific location with formatting
        enhanced_search_and_replace(document_id="main", find_text="EQUATION_HERE", 
                                   replace_with_equation=True, latex_equation=r"rac{a^2 + b^2}{c}",
                                   apply_formatting=True, bold=True)
        
        # Replace text while preserving EndNote citations
        enhanced_search_and_replace(document_id="paper", find_text="old methodology", 
                                   replace_text="new approach", preserve_fields=True)
    """
    from word_document_server.utils.session_utils import resolve_document_path
    
    # Resolve document path from document_id or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    valid_path, filename, path_error = validate_docx_path(filename)
    if not valid_path:
        return path_error

    
    # Validate required parameters
    if not find_text:
        return "Error: find_text parameter is required"
    
    if replace_with_equation:
        if not latex_equation:
            return "Error: latex_equation parameter is required when replace_with_equation is True"
        # Allow replace_text to be None when using replace_with_equation
        # (for pure equation replacement)
    else:
        if replace_text is None:
            return "Error: replace_text parameter is required when replace_with_equation is False"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    size_ok, size_error = check_doc_size_for_operation(filename, "enhanced_search_and_replace")
    if not size_ok:
        return size_error
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    # Validate range / occurrence parameters (Tier-1 & Tier-2)
    if paragraph_indices is not None and (start_paragraph is not None or end_paragraph is not None):
        return "Error: paragraph_indices cannot be combined with start_paragraph/end_paragraph"

    if occurrence_index is not None and occurrence_index < 1:
        return "Error: occurrence_index must be >= 1"

    if (char_start is not None or char_end is not None) and paragraph_indices is None and start_paragraph is None and end_paragraph is None:
        return "Error: char_start/char_end require a paragraph range to be specified"

    if (char_start is None) != (char_end is None):
        return "Error: char_start and char_end must both be provided together"

    if char_start is not None and char_end is not None and char_end < char_start:
        return "Error: char_end cannot be less than char_start"

    # Guardrails
    max_matches_per_call = get_max_matches_per_call()
    max_output_chars = get_max_search_output_chars()
    max_regex_pattern_chars = get_max_regex_pattern_chars()
    max_regex_scan_chars = get_max_regex_scan_chars()
    regex_timeout_ms = get_regex_timeout_ms()

    # Validate regex pattern if using regex
    if use_regex:
        if len(find_text) > max_regex_pattern_chars:
            return (
                f"[{LIMIT_EXCEEDED}] Regex pattern too long. "
                f"len(find_text)={len(find_text)} exceeds EW_MAX_REGEX_PATTERN_CHARS={max_regex_pattern_chars}."
            )

        complexity_reason = _regex_complexity_reason(find_text)
        if complexity_reason:
            return (
                f"[{REGEX_COMPLEXITY_BLOCKED}] Regex pattern refused by preflight: "
                f"{complexity_reason}. Refine the expression."
            )

        if regex_timeout_ms > 0 and not is_regex_timeout_supported():
            return (
                f"[{REGEX_COMPLEXITY_BLOCKED}] EW_REGEX_TIMEOUT_MS={regex_timeout_ms} is configured, "
                "but timeout-capable regex runtime is unavailable. Set EW_REGEX_TIMEOUT_MS=0 "
                "or add runtime support for regex timeouts."
            )

        try:
            import re
            # Validate with the same flags used at runtime
            flags = re.IGNORECASE if not match_case else 0
            re.compile(find_text, flags)
        except re.error as e:
            return f"Invalid regex pattern '{find_text}': {str(e)}"

    # Pre-validate LaTeX conversion once to avoid partial/slow failures during replacement.
    # If conversion fails, we bail out before mutating any document content.
    precomputed_equation_omml_xml = None
    if replace_with_equation:
        ok, omml_or_err = latex_to_omml(latex_equation)
        if not ok:
            return f"Error: latex_equation conversion failed: {omml_or_err}"
        precomputed_equation_omml_xml = omml_or_err
    
    try:
        doc = Document(filename)

        if use_regex:
            regex_scan_chars = sum(len(p.text) for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        regex_scan_chars += sum(len(p.text) for p in cell.paragraphs)
            if regex_scan_chars > max_regex_scan_chars:
                return (
                    f"[{LIMIT_EXCEEDED}] Regex search refused due to scan-size guardrail. "
                    f"document_chars={regex_scan_chars} exceeds EW_MAX_REGEX_SCAN_CHARS={max_regex_scan_chars}. "
                    "Refine the query or raise the limit."
                )
        
        # Determine which paragraphs to process.
        #
        # Note: A valid .docx can contain tables but zero "body paragraphs" as
        # exposed by python-docx. In that case, we still want to process table
        # cell paragraphs when the caller did not specify a paragraph range.
        total_paragraphs = len(doc.paragraphs)
        if total_paragraphs == 0:
            if paragraph_indices is not None or start_paragraph is not None or end_paragraph is not None:
                return "Invalid paragraph range. Document has 0 paragraphs."
            target_paragraphs = []
        elif paragraph_indices is not None:
            target_paragraphs = sorted(set(paragraph_indices))
            if any(idx < 0 or idx >= total_paragraphs for idx in target_paragraphs):
                return f"Invalid paragraph range. Document has {total_paragraphs} paragraphs (0-{total_paragraphs-1})"
        else:
            start_idx = start_paragraph if start_paragraph is not None else 0
            end_idx = end_paragraph if end_paragraph is not None else total_paragraphs - 1
            if start_idx < 0 or end_idx >= total_paragraphs or start_idx > end_idx:
                return f"Invalid paragraph range. Document has {total_paragraphs} paragraphs (0-{total_paragraphs-1})"
            target_paragraphs = list(range(start_idx, end_idx + 1))

        # Global occurrence counter
        current_occurrence = 0
        replacement_limit_hit = False

        def process_para(p_idx, paragraph_list):
            nonlocal current_occurrence, count, equations_inserted, replacement_limit_hit
            # If p_idx is provided, honor paragraph-range gating; for table
            # paragraphs (p_idx is None), process unconditionally.
            if replacement_limit_hit:
                return
            if p_idx is not None and p_idx not in target_paragraphs:
                return
            import re
            para_text_combined = "\n".join([p.text for p in paragraph_list])
            if use_regex:
                pattern_tmp = find_text
                if whole_words_only:
                    pattern_tmp = rf"\b(?:{pattern_tmp})\b"
                flags_tmp = re.IGNORECASE if not match_case else 0
            else:
                escaped = re.escape(find_text)
                pattern_tmp = rf"\b{escaped}\b" if whole_words_only else escaped
                flags_tmp = re.IGNORECASE if not match_case else 0
            try:
                total_matches_para = len(list(re.finditer(pattern_tmp, para_text_combined, flags_tmp)))
            except re.error:
                total_matches_para = 0

            # Determine if this paragraph contains the target occurrence
            local_target = None
            if occurrence_index is not None:
                if current_occurrence + total_matches_para < occurrence_index:
                    # target occurs later; just advance counter
                    current_occurrence += total_matches_para
                    return
                else:
                    local_target = occurrence_index - current_occurrence

            replace_count, match_count, had_equations = _enhanced_replace_in_paragraphs(
                paragraph_list,
                find_text,
                replace_text,
                apply_formatting,
                bold,
                italic,
                underline,
                color,
                font_size,
                font_name,
                match_case,
                whole_words_only,
                use_regex,
                target_occurrence=local_target,
                char_start=char_start,
                char_end=char_end,
                replace_with_equation=replace_with_equation,
                latex_equation=latex_equation,
                preserve_fields=preserve_fields,
                equation_omml_xml=precomputed_equation_omml_xml,
                max_replacements=max_matches_per_call - count,
            )
            if had_equations:
                equations_inserted = True
            current_occurrence += match_count
            count += replace_count
            if count >= max_matches_per_call:
                replacement_limit_hit = True
    
        count = 0
        equations_inserted = False
        

        
        # Process body paragraphs
        for idx, para in enumerate(doc.paragraphs):
            process_para(idx, [para])
            if replacement_limit_hit:
                break

        # Search in tables (treat cell paragraphs independently of doc paragraph indices)
        if not replacement_limit_hit:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        process_para(None, cell.paragraphs)
                        if replacement_limit_hit:
                            break
                    if replacement_limit_hit:
                        break
                if replacement_limit_hit:
                    break
        
        if count > 0:
            # If equations were inserted, ensure math namespace is present
            if equations_inserted:
                from docx.oxml.ns import qn
                root = doc._element
                if "m" not in root.nsmap:
                    root.set(qn("xmlns:m"), "http://schemas.openxmlformats.org/officeDocument/2006/math")
            
            doc.save(filename)
            search_type = "regex pattern" if use_regex else "text"
            case_info = "" if match_case else " (case-insensitive)"
            word_info = " (whole words only)" if whole_words_only else ""
            formatting_applied = " with formatting" if apply_formatting else ""
            
            if replace_with_equation:
                response = (
                    f"Replaced {count} occurrence(s) of {search_type} '{find_text}'"
                    f"{case_info}{word_info} with equation{formatting_applied}."
                )
            else:
                response = (
                    f"Replaced {count} occurrence(s) of {search_type} '{find_text}'"
                    f"{case_info}{word_info} with '{replace_text}'{formatting_applied}."
                )
            if replacement_limit_hit:
                response += (
                    f" [{LIMIT_EXCEEDED}] Replacement limit reached: EW_MAX_MATCHES_PER_CALL={max_matches_per_call}. "
                    "Additional matches were not modified."
                )
            return _truncate_message(response, max_output_chars)
        else:
            search_type = "regex pattern" if use_regex else "text"
            response = f"No occurrences of {search_type} '{find_text}' found."
            return _truncate_message(response, max_output_chars)
            
    except FileNotFoundError:
        return f"Document {filename} not found"
    except PermissionError:
        return f"Permission denied accessing {filename}"
    except Exception as e:
        return f"Failed to search and replace: {str(e)}"


def _enhanced_replace_in_paragraphs(paragraphs, find_text, replace_text, apply_formatting,
                                   bold, italic, underline, color, font_size, font_name,
                                   match_case, whole_words_only, use_regex=False,
                                   target_occurrence=None,
                                   char_start=None, char_end=None,
                                   replace_with_equation=False, latex_equation=None,
                                   preserve_fields=True,
                                   equation_omml_xml=None,
                                   max_replacements=None):
    """Helper function to replace text in paragraphs with optional formatting and regex support.
    
    This implementation fixes the positioning bugs by properly inserting runs at their
    correct positions instead of always appending to the end of the paragraph.
    Supports both literal text matching and regex pattern matching.
    """
    import re
    
    count = 0
    match_total_overall = 0  # total matches found (ignores char-range / occurrence filters)
    equations_inserted = False  # Track if any equations were inserted
    
    for para in paragraphs:
        para_text = para.text
        
        # Create search pattern based on options
        if use_regex:
            # Respect whole_words_only even in regex mode by wrapping pattern
            pattern = find_text
            if whole_words_only:
                pattern = rf"\b(?:{pattern})\b"

            flags = re.IGNORECASE if not match_case else 0
        else:
            # Escape special regex characters for literal matching
            escaped_text = re.escape(find_text)
            if whole_words_only:
                pattern = r'\b' + escaped_text + r'\b'
            else:
                pattern = escaped_text
            flags = re.IGNORECASE if not match_case else 0
        
        # Find all matches in the paragraph text
        try:
            matches = list(re.finditer(pattern, para_text, flags))
        except re.error:
            continue  # Skip if pattern compilation fails
        para_match_total = len(matches)
        match_total_overall += para_match_total

        if not matches:
            continue
            
        matches_to_apply = []
        occurrence_counter = 0  # counts matches that pass char-range gating

        for idx, m in enumerate(matches, start=1):
            # Character-range gating
            if char_start is not None and char_end is not None:
                if not (char_start <= m.start() and m.end() <= char_end):
                    continue  # do NOT count this occurrence toward counter
            # Increment occurrence counter *after* char-range gate
            occurrence_counter += 1

            # Occurrence-specific gating
            if target_occurrence is not None and occurrence_counter != target_occurrence:
                continue

            matches_to_apply.append(m)

            # If we were targeting a single specific occurrence, we can stop after grabbing it
            if target_occurrence is not None:
                break

        if not matches_to_apply:
            continue

        if max_replacements is not None:
            remaining = max_replacements - count
            if remaining <= 0:
                break
            if len(matches_to_apply) > remaining:
                matches_to_apply = matches_to_apply[:remaining]

        count += len(matches_to_apply)
        
        # Process matches from right to left to avoid position shifting during replacement
        for match in reversed(matches_to_apply):
            start_pos = match.start()
            end_pos = match.end()
            matched_text_for_fallback = para_text[start_pos:end_pos]
            
            # For regex, translate JS back-refs then expand groups
            if use_regex and not replace_with_equation:
                converted_repl = _convert_js_backreferences(replace_text)
                actual_replace_text = match.expand(converted_repl)
            elif not replace_with_equation:
                actual_replace_text = replace_text
            else:
                # For equation replacement, we don't need replace_text
                actual_replace_text = None
            
            # For case-insensitive matching, preserve the original case pattern
            if not match_case and not replace_with_equation:
                matched_text = para_text[start_pos:end_pos]
                actual_replace_text = _preserve_case(matched_text, actual_replace_text)
                
            # Handle space collapsing when deleting text
            if not replace_with_equation and replace_text == "":
                # Check if we have spaces before and after the match
                has_space_before = start_pos > 0 and para_text[start_pos - 1].isspace()
                has_space_after = end_pos < len(para_text) and para_text[end_pos].isspace()
                
                # If we have spaces on both sides, keep one space
                if has_space_before and has_space_after:
                    actual_replace_text = " "
            
            # NEW APPROACH: Instead of modifying existing runs and appending new ones,
            # we rebuild the runs in the correct order by collecting all run segments
            # and then reconstructing the paragraph properly.
            
            # Collect all run segments with their formatting and positions
            run_segments = []
            current_pos = 0
            
            for run in para.runs:
                run_length = len(run.text)
                run_start = current_pos
                run_end = current_pos + run_length
                
                # Determine how this run overlaps with the match
                if run_end <= start_pos:
                    # Run is completely before the match - keep as is
                    run_info = format_run_with_citation_awareness(run)
                    if run.text or run_info.get('fields'):  # Keep runs with text or fields
                        run_segments.append({
                            'text': run.text,
                            'formatting': _extract_run_formatting(run),
                            'type': 'keep',
                            'fields': run_info.get('fields'),
                            'run_element': run._element  # Keep reference to original element
                        })
                elif run_start >= end_pos:
                    # Run is completely after the match - keep as is
                    run_info = format_run_with_citation_awareness(run)
                    if run.text or run_info.get('fields'):  # Keep runs with text or fields
                        run_segments.append({
                            'text': run.text,
                            'formatting': _extract_run_formatting(run),
                            'type': 'keep',
                            'fields': run_info.get('fields'),
                            'run_element': run._element  # Keep reference to original element
                        })
                else:
                    # Run overlaps with the match - need to split it
                    
                    # Part before the match
                    if run_start < start_pos:
                        before_text = run.text[:start_pos - run_start]
                        run_info = format_run_with_citation_awareness(run)
                        if before_text:
                            run_segments.append({
                                'text': before_text,
                                'formatting': _extract_run_formatting(run),
                                'type': 'keep',
                                'fields': run_info.get('fields') if run_start == 0 else None  # Only preserve fields if at start
                            })
                    
                    # The match replacement (only add once, when we encounter the first overlapping run)
                    if not any(seg.get('type') == 'replacement' for seg in run_segments):
                        replacement_formatting = _extract_run_formatting(run)
                        if apply_formatting:
                            # Apply new formatting on top of existing
                            replacement_formatting.update({
                                'bold': bold if bold is not None else replacement_formatting.get('bold'),
                                'italic': italic if italic is not None else replacement_formatting.get('italic'),
                                'underline': underline if underline is not None else replacement_formatting.get('underline'),
                                'color': color if color else replacement_formatting.get('color'),
                                'font_size': font_size if font_size else replacement_formatting.get('font_size'),
                                'font_name': font_name if font_name else replacement_formatting.get('font_name')
                            })
                        
                        if replace_with_equation:
                            # Equation OMML should have been prevalidated upstream.
                            if not equation_omml_xml:
                                # Fall back to reinserting original matched text (no-op semantics).
                                run_segments.append({
                                    'text': matched_text_for_fallback,
                                    'formatting': replacement_formatting,
                                    'type': 'replacement'
                                })
                            else:
                                run_segments.append({
                                    'omml_xml': equation_omml_xml,
                                    'formatting': replacement_formatting,
                                    'type': 'equation',
                                    'display_mode': 'inline'  # Default to inline for direct replacement
                                })
                                equations_inserted = True
                        else:
                            # Parse for equation markers in replacement text
                            content_segments = _parse_equation_markers(actual_replace_text)
                            
                            # If only one text segment, add it as before
                            if len(content_segments) == 1 and content_segments[0]['type'] == 'text':
                                run_segments.append({
                                    'text': actual_replace_text,
                                    'formatting': replacement_formatting,
                                    'type': 'replacement'
                                })
                            else:
                                # Multiple segments - add each with appropriate type
                                for content_seg in content_segments:
                                    if content_seg['type'] == 'text' and content_seg['content']:
                                        run_segments.append({
                                            'text': content_seg['content'],
                                            'formatting': replacement_formatting,
                                            'type': 'replacement'
                                        })
                                    elif content_seg['type'] == 'equation':
                                        # Convert LaTeX to OMML
                                        success, omml_or_err = latex_to_omml(content_seg['content'])
                                        if success:
                                            run_segments.append({
                                                'omml_xml': omml_or_err,
                                                'formatting': replacement_formatting,
                                                'type': 'equation',
                                                'display_mode': content_seg.get('display_mode', 'inline')
                                            })
                                            equations_inserted = True
                                        else:
                                            # If conversion fails, add as text
                                            run_segments.append({
                                                'text': f"{{{{equation:{content_seg['content']}}}}}",
                                                'formatting': replacement_formatting,
                                                'type': 'replacement'
                                            })
                    
                    # Part after the match
                    if run_end > end_pos:
                        after_text = run.text[end_pos - run_start:]
                        run_info = format_run_with_citation_awareness(run)
                        if after_text:
                            run_segments.append({
                                'text': after_text,
                                'formatting': _extract_run_formatting(run),
                                'type': 'keep',
                                'fields': run_info.get('fields') if end_pos == run_end else None  # Only preserve fields if at end
                            })
                
                current_pos += run_length
            
            # Clear all existing runs
            for _ in range(len(para.runs)):
                para.runs[0]._element.getparent().remove(para.runs[0]._element)
            
            # Rebuild the paragraph with the correct run segments in order
            # Collapse duplicate spaces across segment boundaries to avoid
            # "double-space" artifacts when replacement text is an empty string.
            cleaned_segments = []
            for seg in run_segments:
                # Keep equation segments regardless of text content
                if seg.get('type') == 'equation':
                    cleaned_segments.append(seg)
                    continue
                # Skip segments with empty text (already handled by check below)
                if seg.get('text', '') == "":
                    continue
                if cleaned_segments:
                    prev = cleaned_segments[-1]
                    # Only process spacing for text segments
                    if 'text' in prev and 'text' in seg:
                        # If previous ends with space and current starts with space → trim one
                        if prev['text'].endswith(' ') and seg['text'].startswith(' '):
                            # Prefer to strip the leading spaces of the current segment to keep formatting of previous run unchanged
                            seg['text'] = seg['text'].lstrip()
                            if seg['text'] == "":
                                # Entire segment became empty → skip it
                                continue
                    # NEW: If previous ends with space and current begins with punctuation, trim the trailing space.
                    punctuation_chars = '.!,?:;)'  # extend as needed
                    if 'text' in prev and 'text' in seg and prev['text'].endswith(' ') and seg['text'][0] in punctuation_chars:
                        prev['text'] = prev['text'].rstrip()
                cleaned_segments.append(seg)

            for segment in cleaned_segments:
                if segment.get('type') == 'equation':
                    # Create empty run for equation
                    # Note: Display mode equations should ideally be in separate paragraphs
                    # Currently all equations are inserted inline within the paragraph
                    new_run = para.add_run()
                    # Parse and append OMML
                    omml_elem = parse_xml(segment['omml_xml'])
                    new_run._r.append(omml_elem)
                    _apply_run_formatting(new_run, segment['formatting'])
                elif segment.get('fields') and len(segment.get('fields', [])) > 0:
                    # This segment contains citation fields
                    # We need to preserve the field structure
                    new_run = para.add_run()
                    # Copy field elements from original run
                    run_element = new_run._element
                    
                    # For simple fields, we need to recreate the fldSimple element
                    for field in segment['fields']:
                        if 'xml_element' in field:
                            # Parse the field XML and append to the new run
                            field_elem = parse_xml(field['xml_element'])
                            run_element.append(field_elem)
                    
                    # Apply formatting to the run
                    _apply_run_formatting(new_run, segment['formatting'])
                elif segment.get('text'):
                    new_run = para.add_run(segment['text'])
                    _apply_run_formatting(new_run, segment['formatting'])
    
    return count, match_total_overall, equations_inserted


def _extract_run_formatting(run):
    """Extract formatting properties from a run."""
    formatting = {}
    try:
        formatting['bold'] = run.bold
        formatting['italic'] = run.italic
        formatting['underline'] = run.underline
        if run.font.name:
            formatting['font_name'] = run.font.name
        if run.font.size:
            formatting['font_size'] = run.font.size.pt if run.font.size else None
        if run.font.color.rgb:
            formatting['color'] = str(run.font.color.rgb)
    except Exception:
        # If any formatting extraction fails, return basic formatting
        pass
    return formatting


def _apply_run_formatting(run, formatting):
    """Apply formatting properties to a run."""
    try:
        if 'bold' in formatting and formatting['bold'] is not None:
            run.bold = formatting['bold']
        if 'italic' in formatting and formatting['italic'] is not None:
            run.italic = formatting['italic']
        if 'underline' in formatting and formatting['underline'] is not None:
            run.underline = formatting['underline']
        if 'font_name' in formatting and formatting['font_name']:
            run.font.name = formatting['font_name']
        if 'font_size' in formatting and formatting['font_size']:
            from docx.shared import Pt
            run.font.size = Pt(formatting['font_size'])
        if 'color' in formatting and formatting['color']:
            _apply_color_to_run(run, formatting['color'])
    except Exception:
        # Silently continue if formatting fails
        pass


def _apply_color_to_run(run, color):
    """Apply color to a run with error handling."""
    from docx.shared import RGBColor
    
    # Define common RGB colors
    color_map = {
        'red': RGBColor(255, 0, 0),
        'blue': RGBColor(0, 0, 255),
        'green': RGBColor(0, 128, 0),
        'yellow': RGBColor(255, 255, 0),
        'black': RGBColor(0, 0, 0),
        'gray': RGBColor(128, 128, 128),
        'white': RGBColor(255, 255, 255),
        'purple': RGBColor(128, 0, 128),
        'orange': RGBColor(255, 165, 0),
        'brown': RGBColor(165, 42, 42),
        'pink': RGBColor(255, 192, 203),
        'cyan': RGBColor(0, 255, 255),
        'magenta': RGBColor(255, 0, 255),
        'lime': RGBColor(0, 255, 0),
        'navy': RGBColor(0, 0, 128),
        'maroon': RGBColor(128, 0, 0),
        'olive': RGBColor(128, 128, 0),
        'teal': RGBColor(0, 128, 128)
    }
    
    try:
        if color.lower() in color_map:
            run.font.color.rgb = color_map[color.lower()]
        else:
            # Accept 3- or 6-digit hex codes, with or without '#', any case
            hex_code = color.lstrip('#')
            if len(hex_code) == 3:  # expand shorthand #FFF → #FFFFFF
                hex_code = ''.join(ch * 2 for ch in hex_code)

            if re.fullmatch(r'[0-9A-Fa-f]{6}', hex_code):
                r = int(hex_code[0:2], 16)
                g = int(hex_code[2:4], 16)
                b = int(hex_code[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            else:
                # Default to black if color not recognized
                run.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        # Fallback to black
        run.font.color.rgb = RGBColor(0, 0, 0)


def format_specific_words(filename: str, word_list: List[str], 
                                bold: Optional[bool] = None,
                                italic: Optional[bool] = None,
                                underline: Optional[bool] = None,
                                color: Optional[str] = None,
                                font_size: Optional[int] = None,
                                font_name: Optional[str] = None,
                                match_case: bool = True,
                                whole_words_only: bool = True,
                                **extra_kwargs) -> str:
    """Format specific words throughout the document using enhanced search and replace.
    
    Args:
        filename: Path to the Word document (resolved path from session management)
        word_list: List of words to format
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        underline: Set text underlined (True/False)
        color: Text color (e.g., 'red', 'blue', etc.)
        font_size: Font size in points
        font_name: Font name/family
        match_case: Whether to match case (default True)
        whole_words_only: Whether to match whole words only (default True)
    """
    results = []
    
    for word in word_list:
        # Use enhanced search and replace with same text for find and replace
        # Pass filename directly since it's already resolved from session management
        result = enhanced_search_and_replace(
            filename=filename,  # Already resolved filename
            find_text=word,
            replace_text=word,  # Same text, just apply formatting
            apply_formatting=True,
            bold=bold,
            italic=italic,
            underline=underline,
            color=color,
            font_size=font_size,
            font_name=font_name,
            match_case=match_case,
            whole_words_only=whole_words_only,
            **extra_kwargs
        )
        results.append(f"'{word}': {result}")
    
    return "\n".join(results)


def format_research_paper_terms(filename: str) -> str:
    """Format common research terms in a PCL paper with appropriate styling - Academic research helper."""
    
    results = []
    
    # Format drug names in blue and bold
    drug_names = ["dolutegravir", "meloxicam", "dexamethasone", "DTG", "MLX", "DEX"]
    result = format_specific_words(filename, drug_names, bold=True, color="blue")
    results.append("Drug names: " + result)
    
    # Format polymer terms in green
    polymer_terms = ["polycaprolactone", "PCL", "mesophase", "crystallinity"]
    result = format_specific_words(filename, polymer_terms, color="green")
    results.append("Polymer terms: " + result)
    
    # Format statistical terms in red and italic
    stats_terms = ["p < 0.05", "significant", "correlation", "ANOVA"]
    result = format_specific_words(filename, stats_terms, italic=True, color="red")
    results.append("Statistical terms: " + result)
    
    # Format temperature values in orange
    temp_terms = ["25°C", "50°C"]
    result = format_specific_words(filename, temp_terms, color="orange")
    results.append("Temperature values: " + result)
    
    return "Research paper terms formatted:\n" + "\n".join(results)



def format_document(
    action: str,
    filename: str = None,
    document_id: str = None,
    word_list: Optional[List[str]] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    color: Optional[str] = None,
    font_size: Optional[int] = None,
    font_name: Optional[str] = None,
    match_case: bool = True,
    whole_words_only: bool = True,
    **extra_kwargs
) -> str:
    """Unified document formatting function for specialized formatting operations.
    
    This consolidated tool replaces 2 individual formatting functions with a single
    action-based interface, reducing tool count while preserving 100% functionality.
    
    Args:
        action (str): Formatting operation to perform:
            - "words": Format specific words in document (requires word_list)
            - "research": Apply research paper formatting (automatic terms)
        filename (str): Path to Word document (legacy, for backward compatibility)
        document_id (str): Session document ID (preferred)
        word_list (List[str], optional): List of words to format (required for "words" action)
        bold (bool, optional): Set text bold (True/False)
        italic (bool, optional): Set text italic (True/False)
        underline (bool, optional): Set text underlined (True/False)
        color (str, optional): Text color (e.g., 'red', 'blue', etc.)
        font_size (int, optional): Font size in points
        font_name (str, optional): Font name/family
        match_case (bool): Whether to match case (default True)
        whole_words_only (bool): Whether to match whole words only (default True)
        
    Returns:
        str: Formatting operation result message
        
    Examples:
        # Format specific words with custom styling
        format_document("words", document_id="thesis", 
                       word_list=["important", "critical", "significant"],
                       bold=True, color="red")
        
        # Apply research paper formatting
        format_document("research", document_id="research_paper")
        
        # Format technical terms
        format_document("words", document_id="manual",
                       word_list=["API", "SDK", "REST"],
                       font_name="Courier New", font_size=11)
    """
    from word_document_server.utils.session_utils import resolve_document_path
    
    # Resolve document path from document_id or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    # Validate action parameter
    valid_actions = ["words", "research"]
    if action not in valid_actions:
        return f"Invalid action: {action}. Must be one of: {', '.join(valid_actions)}"
    
    # Validate required parameters for each action
    if action == "words" and not word_list:
        return "Error: 'word_list' is required for action 'words'"
    
    # Delegate to appropriate original function based on action
    try:
        if action == "words":
            return format_specific_words(
                filename=filename,
                word_list=word_list,
                bold=bold,
                italic=italic,
                underline=underline,
                color=color,
                font_size=font_size,
                font_name=font_name,
                match_case=match_case,
                whole_words_only=whole_words_only,
                **extra_kwargs
            )
            
        elif action == "research":
            return format_research_paper_terms(filename)
            
    except Exception as e:
        return f"Error in format_document: {str(e)}"


# ---------------------------------------------------------------------------
# Helper: convert JavaScript-style regex replacement markers ($&,$1,$2, ...) to
# Python `re` syntax (\g<0>, \g<1>, ...). This allows end-users coming from
# JS/VS Code to keep familiar notation while leveraging Python's engine.
# ---------------------------------------------------------------------------


def _convert_js_backreferences(repl: str) -> str:
    """Translate $& and $1…$99 tokens to Python replacement notation.

    Example:
        "$&"  -> "\\g<0>"
        "$1_$2" -> "\\g<1>_\\g<2>"
    """

    # Replace $& (whole match) first to avoid interference with $1..
    repl = repl.replace("$&", r"\g<0>")

    # Replace numbered groups – supports up to two-digit back-refs
    def _num_sub(m):
        return rf"\g<{m.group(1)}>"

    return re.sub(r"\$(\d{1,2})", _num_sub, repl)

def _preserve_case(original_text: str, replacement_text: str) -> str:
    """Preserve the case pattern from original text in the replacement text.
    
    Args:
        original_text: The original text whose case pattern should be preserved
        replacement_text: The replacement text to apply the case pattern to
        
    Returns:
        The replacement text with the case pattern from the original text
    """
    if not original_text or not replacement_text:
        return replacement_text
    
    # Check case patterns
    if original_text.isupper():
        # All uppercase
        return replacement_text.upper()
    elif original_text.islower():
        # All lowercase
        return replacement_text.lower()
    elif original_text[0].isupper() and original_text[1:].islower():
        # Title case
        return replacement_text.capitalize()
    else:
        # Mixed case - preserve character-by-character where possible
        result = []
        for i, char in enumerate(replacement_text):
            if i < len(original_text):
                if original_text[i].isupper():
                    result.append(char.upper())
                else:
                    result.append(char.lower())
            else:
                # Beyond original length, keep replacement as-is
                result.append(char)
        return ''.join(result)

def _parse_equation_markers(text: str) -> List[dict]:
    """Parse text containing {{equation:latex}} markers into segments.
    
    Supports formats:
    - {{equation:latex}} - Auto-detect display/inline based on context
    - {{equation:display:latex}} - Force display mode
    - {{equation:inline:latex}} - Force inline mode
    - {{equation1:latex}}, {{equation2:latex}} - Numbered equations
    
    Returns a list of segments with either 'text' or 'equation' type.
    """
    import re
    segments = []
    last_end = 0
    
    # Enhanced pattern to capture optional display mode and equation numbers
    # Matches: {{equation:latex}}, {{equation:display:latex}}, {{equation1:latex}}, etc.
    pattern = r'\{\{equation(\d*):(?:(display|inline):)?(.*?)\}\}'
    
    for match in re.finditer(pattern, text):
        # Add text before the equation marker
        if match.start() > last_end:
            text_before = text[last_end:match.start()]
            if text_before:
                segments.append({'type': 'text', 'content': text_before})
        
        # Extract equation details
        eq_number = match.group(1) or ''
        display_mode = match.group(2)  # 'display', 'inline', or None for auto
        latex = match.group(3).strip()
        
        # Auto-detect display mode if not specified
        if display_mode is None:
            # Check context for auto-detection
            text_before = text[last_end:match.start()].strip()
            text_after_end = match.end()
            text_after = text[text_after_end:text_after_end + 50].strip() if text_after_end < len(text) else ''
            
            # If equation is alone in the paragraph or at start/end, make it display
            if (not text_before or text_before.endswith('\n')) and (not text_after or text_after.startswith('\n')):
                display_mode = 'display'
            else:
                display_mode = 'inline'
        
        segments.append({
            'type': 'equation',
            'content': latex,
            'display_mode': display_mode,
            'number': eq_number
        })
        
        last_end = match.end()
    
    # Add any remaining text after the last equation
    if last_end < len(text):
        remaining_text = text[last_end:]
        if remaining_text:
            segments.append({'type': 'text', 'content': remaining_text})
    
    # If no equations found, return single text segment
    if not segments:
        segments.append({'type': 'text', 'content': text})
    
    return segments
