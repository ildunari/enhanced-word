from __future__ import annotations

import asyncio
from pathlib import Path

from docx import Document

from word_document_server.tools.content_tools import enhanced_search_and_replace
from word_document_server.tools.document_tools import create_document


def test_enhanced_search_and_replace_can_match_across_runs(tmp_path: Path):
    p = tmp_path / "runs.docx"
    asyncio.run(create_document(filename=str(p)))

    doc = Document(str(p))
    para = doc.add_paragraph()
    para.add_run("Hello")
    para.add_run("World")
    doc.save(str(p))

    res = enhanced_search_and_replace(filename=str(p), find_text="HelloWorld", replace_text="Hi")
    assert "replaced" in res.lower()

    doc2 = Document(str(p))
    assert "Hi" in "\n".join([para.text for para in doc2.paragraphs])


def test_enhanced_search_and_replace_updates_table_cells(tmp_path: Path):
    p = tmp_path / "table.docx"
    asyncio.run(create_document(filename=str(p)))

    doc = Document(str(p))
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "foo"
    doc.save(str(p))

    res = enhanced_search_and_replace(filename=str(p), find_text="foo", replace_text="bar")
    assert "replaced" in res.lower()

    doc2 = Document(str(p))
    assert doc2.tables[0].cell(0, 0).text == "bar"


def test_enhanced_search_and_replace_deleting_word_collapses_spaces(tmp_path: Path):
    p = tmp_path / "spaces.docx"
    asyncio.run(create_document(filename=str(p)))

    doc = Document(str(p))
    doc.add_paragraph("a X b")
    doc.save(str(p))

    res = enhanced_search_and_replace(
        filename=str(p),
        find_text="X",
        replace_text="",
        whole_words_only=True,
    )
    assert "replaced" in res.lower()

    doc2 = Document(str(p))
    assert "a b" in "\n".join([para.text for para in doc2.paragraphs])

