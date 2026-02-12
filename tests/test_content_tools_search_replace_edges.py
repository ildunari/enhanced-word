from __future__ import annotations

import asyncio
from pathlib import Path

from docx import Document

from word_document_server.tools import content_tools
from word_document_server.tools.content_tools import add_text_content, enhanced_search_and_replace
from word_document_server.tools.document_tools import create_document


def test_enhanced_search_and_replace_regex_backrefs(tmp_path: Path):
    p = tmp_path / "regex.docx"
    asyncio.run(create_document(filename=str(p)))
    asyncio.run(add_text_content(filename=str(p), text="Date: 2026-02-12"))

    res = enhanced_search_and_replace(
        filename=str(p),
        find_text=r"(\d{4})-(\d{2})-(\d{2})",
        replace_text="$2/$3/$1",
        use_regex=True,
    )
    assert "replaced" in res.lower() or "no occurrences" in res.lower()

    doc = Document(str(p))
    assert "02/12/2026" in "\n".join([para.text for para in doc.paragraphs])


def test_enhanced_search_and_replace_case_insensitive_preserves_case(tmp_path: Path):
    p = tmp_path / "case.docx"
    asyncio.run(create_document(filename=str(p)))
    asyncio.run(add_text_content(filename=str(p), text="FOO foo Foo"))

    res = enhanced_search_and_replace(
        filename=str(p),
        find_text="foo",
        replace_text="bar",
        match_case=False,
    )
    assert "replaced" in res.lower()

    doc = Document(str(p))
    assert doc.paragraphs[0].text == "BAR bar Bar"


def test_enhanced_search_and_replace_occurrence_index_replaces_one(tmp_path: Path):
    p = tmp_path / "occ.docx"
    asyncio.run(create_document(filename=str(p)))
    asyncio.run(add_text_content(filename=str(p), text="x x x"))

    res = enhanced_search_and_replace(
        filename=str(p),
        find_text="x",
        replace_text="y",
        whole_words_only=True,
        occurrence_index=2,
    )
    assert "replaced" in res.lower()

    doc = Document(str(p))
    assert doc.paragraphs[0].text == "x y x"


def test_replace_with_equation_prevalidates_once_and_inserts_omml(tmp_path: Path, monkeypatch):
    p = tmp_path / "eq.docx"
    asyncio.run(create_document(filename=str(p)))
    asyncio.run(add_text_content(filename=str(p), text="EQ"))

    calls: list[str] = []

    def _fake_latex_to_omml(latex: str):
        calls.append(latex)
        # Minimal OMML container; needs the m namespace to parse.
        omml = (
            "<m:oMath xmlns:m=\"http://schemas.openxmlformats.org/officeDocument/2006/math\">"
            "<m:r><m:t>1</m:t></m:r>"
            "</m:oMath>"
        )
        return True, omml

    monkeypatch.setattr(content_tools, "latex_to_omml", _fake_latex_to_omml)

    res = enhanced_search_and_replace(
        filename=str(p),
        find_text="EQ",
        replace_text=None,
        replace_with_equation=True,
        latex_equation=r"E = mc^2",
    )
    assert "with equation" in res.lower()
    assert calls == [r"E = mc^2"]

    doc = Document(str(p))
    # python-docx doesn't expose equations directly; validate via XML.
    assert "oMath" in doc.paragraphs[0]._p.xml


def test_replace_with_equation_invalid_latex_returns_error_and_no_mutation(tmp_path: Path, monkeypatch):
    p = tmp_path / "eqbad.docx"
    asyncio.run(create_document(filename=str(p)))
    asyncio.run(add_text_content(filename=str(p), text="EQ"))

    def _bad_latex_to_omml(_latex: str):
        return False, "bad latex"

    monkeypatch.setattr(content_tools, "latex_to_omml", _bad_latex_to_omml)

    res = enhanced_search_and_replace(
        filename=str(p),
        find_text="EQ",
        replace_text=None,
        replace_with_equation=True,
        latex_equation="not valid",
    )
    assert res.lower().startswith("error:")

    doc = Document(str(p))
    assert doc.paragraphs[0].text.strip() == "EQ"


def test_enhanced_search_and_replace_requires_both_char_bounds(tmp_path: Path):
    p = tmp_path / "char-range.docx"
    asyncio.run(create_document(filename=str(p)))
    asyncio.run(add_text_content(filename=str(p), text="foo bar foo"))

    res = enhanced_search_and_replace(
        filename=str(p),
        find_text="foo",
        replace_text="X",
        start_paragraph=0,
        end_paragraph=0,
        char_start=4,
    )
    assert "char_start and char_end must both be provided together" in res

    doc = Document(str(p))
    assert doc.paragraphs[0].text == "foo bar foo"
