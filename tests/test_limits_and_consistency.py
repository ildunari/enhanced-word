from __future__ import annotations

import asyncio
from pathlib import Path

from docx import Document

from word_document_server.session_manager import get_session_manager
from word_document_server.tools.content_tools import add_text_content, add_table, enhanced_search_and_replace
from word_document_server.tools.document_tools import get_text
from word_document_server.tools.session_tools import open_document


def _write_large_doc(path: Path, paragraphs: int = 60, width: int = 300) -> None:
    doc = Document()
    for i in range(paragraphs):
        doc.add_paragraph(f"{i:04d} " + ("x" * width))
    doc.save(str(path))


def test_doc_size_guard_blocks_add_text_without_mutating(tmp_path: Path, monkeypatch):
    p = tmp_path / "large-add.docx"
    _write_large_doc(p, paragraphs=40, width=350)

    monkeypatch.setenv("EW_MAX_DOC_BYTES_PER_OPERATION", "4000")
    before = p.read_bytes()
    result = asyncio.run(add_text_content(filename=str(p), text="new line"))

    assert "DOC_TOO_LARGE" in result
    assert p.read_bytes() == before


def test_doc_size_guard_blocks_search_replace_without_mutating(tmp_path: Path, monkeypatch):
    p = tmp_path / "large-replace.docx"
    _write_large_doc(p, paragraphs=45, width=320)

    assert p.stat().st_size > 1000
    monkeypatch.setenv("EW_MAX_DOC_BYTES_PER_OPERATION", "1000")
    before = p.read_bytes()
    result = enhanced_search_and_replace(filename=str(p), find_text="xxxx", replace_text="yyyy")

    assert "DOC_TOO_LARGE" in result
    assert p.read_bytes() == before


def test_get_text_search_refuses_oversized_document(tmp_path: Path, monkeypatch):
    p = tmp_path / "large-search.docx"
    _write_large_doc(p, paragraphs=50, width=280)

    monkeypatch.setenv("EW_MAX_DOC_BYTES_PER_OPERATION", "4000")
    result = asyncio.run(get_text(filename=str(p), scope="search", search_term="x"))
    assert "DOC_TOO_LARGE" in result


def test_filename_and_document_id_paths_produce_consistent_edits(tmp_path: Path):
    p = tmp_path / "consistency.docx"
    doc = Document()
    doc.add_paragraph("alpha beta beta")
    doc.save(str(p))

    assert "successfully opened" in open_document("main", str(p)).lower()

    result_file = enhanced_search_and_replace(
        filename=str(p),
        find_text="beta",
        replace_text="gamma",
        occurrence_index=1,
    )
    assert "replaced" in result_file.lower()

    result_id = enhanced_search_and_replace(
        document_id="main",
        find_text="beta",
        replace_text="delta",
        occurrence_index=1,
    )
    assert "replaced" in result_id.lower()

    text_via_file = asyncio.run(get_text(filename=str(p), scope="all"))
    text_via_id = asyncio.run(get_text(document_id="main", scope="all"))
    assert text_via_file == text_via_id
    assert "alpha gamma delta" in text_via_file


def test_doc_extension_rejection_is_consistent_across_tools():
    add_text_res = asyncio.run(add_text_content(filename="bad.doc", text="x"))
    add_table_res = asyncio.run(add_table(filename="bad.doc", rows=1, cols=1))
    get_text_res = asyncio.run(get_text(filename="bad.doc", scope="all"))
    replace_res = enhanced_search_and_replace(filename="bad.doc", find_text="a", replace_text="b")

    for result in [add_text_res, add_table_res, get_text_res, replace_res]:
        assert ".doc files are not supported" in result


def test_long_session_operation_limit_is_enforced(tmp_path: Path, monkeypatch):
    p = tmp_path / "ops.docx"
    doc = Document()
    doc.add_paragraph("seed")
    doc.save(str(p))

    monkeypatch.setenv("EW_LONG_SESSION_OP_LIMIT", "2")
    mgr = get_session_manager()
    assert "successfully opened" in mgr.open_document("ops", str(p)).lower()

    assert "seed" in asyncio.run(get_text(document_id="ops", scope="all"))
    assert "seed" in asyncio.run(get_text(document_id="ops", scope="all"))

    third = asyncio.run(get_text(document_id="ops", scope="all"))
    assert "SESSION_CONSISTENCY_WARNING" in third
