from __future__ import annotations

import asyncio
from pathlib import Path

import pytest
from docx import Document

from word_document_server.tools.content_tools import add_text_content
from word_document_server.tools.document_tools import get_text, merge_documents
from word_document_server.tools.session_tools import open_document
from word_document_server.tools.undo_tools import session_undo


pytestmark = pytest.mark.stress


def _write_large_doc(path: Path, paragraphs: int = 120, width: int = 340) -> None:
    doc = Document()
    for i in range(paragraphs):
        doc.add_paragraph(f"{i:04d} " + ("x" * width))
    doc.save(str(path))


def test_stress_search_scope_refuses_oversized_document(tmp_path: Path, monkeypatch):
    p = tmp_path / "oversized-search.docx"
    _write_large_doc(p)
    monkeypatch.setenv("EW_MAX_DOC_BYTES_PER_OPERATION", "6000")

    result = asyncio.run(get_text(filename=str(p), scope="search", search_term="x"))
    assert "DOC_TOO_LARGE" in result


def test_stress_merge_refuses_oversized_source(tmp_path: Path, monkeypatch):
    source = tmp_path / "source-large.docx"
    target = tmp_path / "target.docx"
    _write_large_doc(source, paragraphs=100, width=360)
    monkeypatch.setenv("EW_MAX_DOC_BYTES_PER_OPERATION", "7000")

    result = asyncio.run(merge_documents(target_filename=str(target), source_filenames=[str(source)]))
    assert "DOC_TOO_LARGE" in result


def test_stress_undo_budget_evicts_oldest_snapshots(tmp_path: Path, monkeypatch):
    p = tmp_path / "undo-budget.docx"
    doc = Document()
    doc.add_paragraph("seed")
    doc.save(str(p))

    monkeypatch.setenv("EW_MAX_UNDO_BYTES_TOTAL", "30000")
    assert "successfully opened" in open_document("budget", str(p)).lower()

    for i in range(18):
        result = asyncio.run(add_text_content(document_id="budget", text=f"line-{i}-" + ("x" * 500)))
        assert "added" in result.lower()

    history = session_undo(action="list", document_id="budget")
    assert "budget_evictions=" in history
    assert "UNDO_BUDGET_EXCEEDED" in history
