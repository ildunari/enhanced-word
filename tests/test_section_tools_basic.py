from __future__ import annotations

import asyncio
import json
from pathlib import Path

from docx import Document

from word_document_server.tools.section_tools import get_sections


def test_get_sections_overview_json_includes_heading(tmp_path: Path):
    p = tmp_path / "s.docx"
    doc = Document()
    doc.add_heading("Intro", level=1)
    doc.add_paragraph("Body")
    doc.save(str(p))

    res = asyncio.run(get_sections(filename=str(p), mode="overview", output_format="json"))
    data = json.loads(res)

    # Be tolerant to output shape; just ensure heading title appears somewhere.
    blob = json.dumps(data)
    assert "Intro" in blob

