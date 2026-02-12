from __future__ import annotations

from pathlib import Path
import tempfile

from docx import Document
from hypothesis import HealthCheck, given, settings, strategies as st, assume

from word_document_server.tools.content_tools import enhanced_search_and_replace


def _write_single_paragraph(path: Path, text: str) -> None:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(path))


def _write_fragmented_paragraph(path: Path, chunks: list[str]) -> None:
    doc = Document()
    para = doc.add_paragraph()
    for chunk in chunks:
        para.add_run(chunk)
    doc.save(str(path))


@settings(max_examples=30, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
@given(
    source=st.text(alphabet=st.characters(min_codepoint=32, max_codepoint=126), min_size=1, max_size=60),
    needle=st.text(alphabet=st.characters(min_codepoint=33, max_codepoint=126), min_size=1, max_size=6),
)
def test_replace_absent_token_is_noop(source: str, needle: str):
    assume(needle not in source)

    with tempfile.TemporaryDirectory(prefix="ew-prop-noop-") as tmp:
        path = Path(tmp) / "noop.docx"
        _write_single_paragraph(path, source)

        result = enhanced_search_and_replace(filename=str(path), find_text=needle, replace_text="ZZ")
        assert "No occurrences" in result

        after = Document(str(path)).paragraphs[0].text
        assert after == source


@settings(max_examples=30, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
@given(source=st.text(alphabet=st.characters(min_codepoint=32, max_codepoint=126), min_size=1, max_size=80))
def test_deletion_never_increases_length(source: str):
    with tempfile.TemporaryDirectory(prefix="ew-prop-del-") as tmp:
        path = Path(tmp) / "del.docx"
        _write_single_paragraph(path, source)

        _ = enhanced_search_and_replace(filename=str(path), find_text="x", replace_text="")
        after = Document(str(path)).paragraphs[0].text
        assert len(after) <= len(source)


@settings(max_examples=30, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
@given(
    prefix=st.text(alphabet=st.characters(min_codepoint=32, max_codepoint=126), min_size=0, max_size=20),
    middle=st.text(alphabet=st.characters(min_codepoint=32, max_codepoint=126), min_size=0, max_size=20),
    suffix=st.text(alphabet=st.characters(min_codepoint=32, max_codepoint=126), min_size=0, max_size=20),
    chunk_sizes=st.lists(st.integers(min_value=1, max_value=5), min_size=1, max_size=25),
)
def test_fragmented_runs_match_single_run_replacement(prefix: str, middle: str, suffix: str, chunk_sizes: list[int]):
    base = f"{prefix}TOKEN{middle}TOKEN{suffix}"

    # Deterministically fragment `base` according to generated chunk sizes.
    chunks: list[str] = []
    i = 0
    j = 0
    while i < len(base):
        step = chunk_sizes[j % len(chunk_sizes)]
        chunks.append(base[i : i + step])
        i += step
        j += 1

    with tempfile.TemporaryDirectory(prefix="ew-prop-runs-") as tmp:
        single_path = Path(tmp) / "single.docx"
        fragmented_path = Path(tmp) / "fragmented.docx"

        _write_single_paragraph(single_path, base)
        _write_fragmented_paragraph(fragmented_path, chunks)

        _ = enhanced_search_and_replace(filename=str(single_path), find_text="TOKEN", replace_text="X")
        _ = enhanced_search_and_replace(filename=str(fragmented_path), find_text="TOKEN", replace_text="X")

        single_text = Document(str(single_path)).paragraphs[0].text
        fragmented_text = Document(str(fragmented_path)).paragraphs[0].text
        assert fragmented_text == single_text

