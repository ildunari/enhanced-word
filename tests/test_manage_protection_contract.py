from __future__ import annotations

import asyncio
from pathlib import Path

from word_document_server.core.protection import add_protection_info
from word_document_server.tools.protection_tools import manage_protection


def test_manage_protection_restricted_roundtrip_uses_canonical_metadata_path(tmp_path: Path):
    p = tmp_path / "restricted.docx"
    from docx import Document

    Document().save(str(p))
    metadata_path = p.with_suffix(".protection")

    protect_res = asyncio.run(
        manage_protection(
            filename=str(p),
            action="protect",
            protection_type="restricted",
            password="secret",
            editable_sections=["Intro"],
        )
    )
    assert "restricted editing protection added" in protect_res.lower()
    assert metadata_path.exists()

    status_res = asyncio.run(
        manage_protection(filename=str(p), action="status", protection_type="restricted")
    )
    assert "restricted editing protection" in status_res.lower()

    unprotect_res = asyncio.run(
        manage_protection(
            filename=str(p),
            action="unprotect",
            protection_type="restricted",
            password="secret",
        )
    )
    assert "restricted editing protection removed" in unprotect_res.lower()
    assert not metadata_path.exists()


def test_manage_protection_signature_roundtrip_visible_in_status_and_verify(tmp_path: Path):
    from docx import Document

    p = tmp_path / "signature.docx"
    doc = Document()
    doc.add_paragraph("Hello")
    doc.save(str(p))

    protect_res = asyncio.run(
        manage_protection(
            filename=str(p),
            action="protect",
            protection_type="signature",
            signer_name="Alice",
            signature_reason="Approval",
        )
    )
    assert "digital signature added" in protect_res.lower()
    assert p.with_suffix(".protection").exists()

    status_res = asyncio.run(
        manage_protection(filename=str(p), action="status", protection_type="signature")
    )
    assert "digitally signed by alice" in status_res.lower()

    verify_res = asyncio.run(
        manage_protection(filename=str(p), action="verify", protection_type="signature")
    )
    assert "digital signature verified" in verify_res.lower()


def test_manage_protection_password_does_not_report_success_without_verifiable_state(tmp_path: Path):
    from docx import Document

    p = tmp_path / "password.docx"
    doc = Document()
    doc.add_paragraph("Secret")
    doc.save(str(p))

    protect_res = asyncio.run(
        manage_protection(
            filename=str(p),
            action="protect",
            protection_type="password",
            password="abc123",
        )
    )

    if "password protection added" in protect_res.lower():
        status_res = asyncio.run(
            manage_protection(filename=str(p), action="status", protection_type="password")
        )
        assert "not password protected" not in status_res.lower()


def test_manage_protection_status_reads_existing_restricted_metadata(tmp_path: Path):
    from docx import Document

    p = tmp_path / "existing.docx"
    Document().save(str(p))

    ok = add_protection_info(
        str(p),
        protection_type="restricted",
        password_hash="hash",
        sections=["A"],
    )
    assert ok

    status_res = asyncio.run(
        manage_protection(filename=str(p), action="status", protection_type="restricted")
    )
    assert "restricted editing protection" in status_res.lower()
