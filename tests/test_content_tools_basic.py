from __future__ import annotations

import asyncio
from pathlib import Path

from docx import Document

from word_document_server.tools.document_tools import create_document
from word_document_server.tools.content_tools import add_text_content, enhanced_search_and_replace


def test_add_text_content_appends_paragraph(tmp_path: Path):
    p = tmp_path / "c.docx"
    asyncio.run(create_document(filename=str(p)))

    res = asyncio.run(add_text_content(filename=str(p), text="Hello world"))
    assert not res.lower().startswith("error")

    doc = Document(str(p))
    assert doc.paragraphs[-1].text.strip() == "Hello world"


def test_enhanced_search_and_replace_replaces_text(tmp_path: Path):
    p = tmp_path / "r.docx"
    asyncio.run(create_document(filename=str(p)))
    asyncio.run(add_text_content(filename=str(p), text="foo bar"))

    res = enhanced_search_and_replace(filename=str(p), find_text="bar", replace_text="baz")
    assert "replace" in res.lower() or "replaced" in res.lower() or "match" in res.lower()

    doc = Document(str(p))
    full_text = "\n".join([para.text for para in doc.paragraphs])
    assert "foo baz" in full_text
