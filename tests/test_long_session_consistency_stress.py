from __future__ import annotations

import asyncio
import base64
from pathlib import Path

import pytest
from docx import Document

from word_document_server.tools.content_tools import add_picture, add_table, add_text_content, enhanced_search_and_replace
from word_document_server.tools.document_tools import get_text
from word_document_server.tools.session_tools import open_document
from word_document_server.tools.undo_tools import session_undo


pytestmark = pytest.mark.stress


def _write_base_doc(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("session seed")
    doc.save(str(path))


def _write_png(path: Path) -> None:
    path.write_bytes(
        base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO6p1U8AAAAASUVORK5CYII=")
    )


def test_stress_long_session_mixed_operations_remains_consistent(tmp_path: Path):
    doc_path = tmp_path / "long-session.docx"
    img_path = tmp_path / "dot.png"
    _write_base_doc(doc_path)
    _write_png(img_path)

    assert "successfully opened" in open_document("main", str(doc_path)).lower()

    for i in range(120):
        mod = i % 6
        if mod == 0:
            result = asyncio.run(add_text_content(document_id="main", text=f"line-{i}"))
            assert "added" in result.lower()
        elif mod == 1:
            result = enhanced_search_and_replace(
                document_id="main",
                find_text="line",
                replace_text="line",
                whole_words_only=True,
            )
            assert "replaced" in result.lower() or "no occurrences" in result.lower()
        elif mod == 2:
            result = asyncio.run(add_table(document_id="main", rows=1, cols=1, data=[[str(i)]]))
            assert "added" in result.lower()
        elif mod == 3:
            result = asyncio.run(get_text(document_id="main", scope="search", search_term="line", max_results=5))
            assert isinstance(result, str)
        elif mod == 4:
            result = asyncio.run(add_picture(document_id="main", image_path=str(img_path)))
            assert "added" in result.lower()
        else:
            listing = session_undo(action="list", document_id="main")
            assert "History for" in listing or "No undo history tracked yet" in listing

    doc = Document(str(doc_path))
    assert len(doc.paragraphs) >= 1
    assert len(doc.tables) >= 1

    undo_res = session_undo(action="undo", steps=1, document_id="main")
    assert "undo" in undo_res.lower()

    # Must still be a readable .docx after long mixed sessions.
    doc_after = Document(str(doc_path))
    assert len(doc_after.paragraphs) >= 1


def test_stress_path_is_authoritative_over_cached_handle(tmp_path: Path):
    doc_path = tmp_path / "path-authority.docx"
    _write_base_doc(doc_path)
    assert "successfully opened" in open_document("main", str(doc_path)).lower()

    # Mutate on disk directly after opening session document.
    disk_doc = Document(str(doc_path))
    disk_doc.add_paragraph("disk-only-change")
    disk_doc.save(str(doc_path))

    observed = asyncio.run(get_text(document_id="main", scope="all"))
    assert "disk-only-change" in observed
